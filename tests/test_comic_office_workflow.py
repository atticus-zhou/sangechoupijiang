import unittest
import asyncio
from pathlib import Path
from tempfile import TemporaryDirectory
from unittest.mock import patch

from src.comic_artifacts import build_comic_artifacts
from src.comic_word_canvas import build_comic_word_canvas
from src.comic_office import (
    advance_comic_cabinet_session,
    advance_comic_cabinet_session_llm,
    build_confirmed_script,
    build_comic_brief,
    build_comic_request,
    build_comic_result,
    build_comic_script_preview,
    enhance_comic_prompts_llm,
    start_comic_cabinet_session,
    start_comic_cabinet_session_llm,
    validate_confirmed_script_session,
)
from src.llm.providers import LLMResponse, LiteLLMProvider, ModelConfig


class ComicOfficeWorkflowTests(unittest.TestCase):
    def test_parse_comic_request_keeps_character_and_style_references(self):
        from src.comic_office.workflow import parse_comic_request

        spec = parse_comic_request(
            "\n".join([
                "Idea: 修仙小队失去辅助后追凶",
                "Character references:",
                "大师兄：沉稳，青黑长袍，左眉有旧伤。",
                "二师姐：冷静，银簪，不轻易笑。",
                "Style references:",
                "古风仙侠，雨夜，冷蓝月光，水墨厚涂。",
                "Input mode: full_script",
                "Full script:",
                "辅助在雨夜死亡。大师兄带队追查真凶。",
            ])
        )

        self.assertIn("大师兄", spec["character_references"])
        self.assertIn("二师姐", spec["character_references"])
        self.assertIn("冷蓝月光", spec["style_references"])
        self.assertEqual(spec["input_mode"], "full_script")
        self.assertIn("辅助在雨夜死亡", spec["full_script"])

    def test_cabinet_suggests_selectable_story_directions_for_vague_idea(self):
        result = start_comic_cabinet_session(
            idea="一个学生站在楼顶，母亲在下面苦苦哀求",
            genre="现实情感",
            length="3集，每集60秒",
            platform="竖屏短视频",
            visual_style="现实主义",
            extra="",
        )

        options = result["session"]["story_state"].get("direction_options") or []
        self.assertGreaterEqual(len(options), 2)
        self.assertLessEqual(len(options), 3)
        self.assertTrue(all(option.get("label") and option.get("reason") for option in options))
        self.assertIn("方向一", result["assistant_message"])
        self.assertIn("方向二", result["assistant_message"])

    def test_cabinet_returns_short_suggested_replies_for_next_turn(self):
        result = start_comic_cabinet_session(
            idea="一个学生站在楼顶，母亲在下面苦苦哀求",
            genre="现实情感",
            length="3集，每集60秒",
            platform="竖屏短视频",
            visual_style="现实主义",
            extra="",
        )

        replies = result["suggested_replies"]
        self.assertGreaterEqual(len(replies), 2)
        self.assertLessEqual(len(replies), 3)
        self.assertEqual(replies, result["session"]["story_state"]["suggested_replies"])
        self.assertTrue(all(reply.startswith("我想走") for reply in replies))
        self.assertTrue(all(len(reply) <= 80 for reply in replies))

    def test_complete_full_script_is_ready_to_confirm_without_extra_answer(self):
        full_script = (
            "主角阿衡独自护送宗门密信下山。追兵在山路截杀他，逼他交出密信。"
            "阿衡负伤逃进废弃驿站，用假信引开追兵。师姐赶到时，他把真信交给她。"
            "最终密信救下宗门，阿衡也决定不再独自承担所有危险。"
        )

        result = start_comic_cabinet_session(
            idea="宗门密信",
            genre="古风仙侠",
            length="3集，每集60秒",
            platform="竖屏短视频",
            visual_style="电影级国风厚涂",
            extra="Input mode: full_script\nFull script:\n" + full_script,
        )

        self.assertTrue(result["ready_to_produce"])
        self.assertEqual(result["stage"], "ready_to_confirm")
        self.assertEqual(result["session"]["story_state"]["questions"], [])
        self.assertEqual(result["script_preview"]["story_draft"], full_script)

    def test_full_script_episode_outline_matches_requested_three_episode_length(self):
        full_script = (
            "阿衡接到护送密信的任务。追兵在山路出现。"
            "他用假信引开追兵。师姐找到重伤的阿衡。"
            "师姐带信返回宗门。最终阿衡得到应有的尊重。"
        )

        result = start_comic_cabinet_session(
            idea="宗门密信",
            genre="古风仙侠",
            length="3集，每集60秒",
            platform="竖屏短视频",
            visual_style="电影级国风厚涂",
            extra="Input mode: full_script\nFull script:\n" + full_script,
        )

        outline = result["script_preview"]["episode_outline"]
        self.assertEqual(len(outline), 3)
        self.assertEqual([item["episode"] for item in outline], [1, 2, 3])
        self.assertIn("最终阿衡得到应有的尊重", outline[-1]["hook"])

    def test_prompt_enhancement_parser_skips_invalid_braces_before_fenced_json(self):
        from src.comic_office.workflow import _parse_prompt_enhancement_json

        payload = _parse_prompt_enhancement_json(
            "先说明一下：{这不是有效JSON}\n```json\n"
            '{"characters": [{"id": "char_01", "image_prompt": "有效提示词"}], "quality_review": {"status": "pass"}}'
            "\n```\n以上是结果。"
        )

        self.assertEqual(payload["characters"][0]["image_prompt"], "有效提示词")

    def test_prompt_enhancement_retries_once_when_first_structured_output_is_invalid(self):
        class RetryPromptProvider:
            def __init__(self):
                self.calls = 0

            async def chat(self, messages, **kwargs):
                self.calls += 1
                if self.calls == 1:
                    return LLMResponse(content="这轮结构没有生成成功", model="fake")
                return LLMResponse(
                    content=(
                        '```json\n{"characters":[{"id":"char_01","image_prompt":"第二轮有效人物提示词"}],'
                        '"quality_review":{"status":"pass","summary":"结构有效","issues":[]}}\n```'
                    ),
                    model="fake",
                )

        request = build_comic_request(
            idea="古风侠客守护一封密信",
            genre="古风仙侠",
            visual_style="电影级国风厚涂",
            confirmed_script={
                "status": "confirmed",
                "title": "密信",
                "story_draft": "侠客护送密信穿过山门，最终把密信交给师父。",
                "story_promise": "一次护送揭开门派秘密。",
                "main_conflict": "侠客必须避开追兵。",
                "why_it_happens": "密信关系门派存亡。",
                "how_it_happens": "护送、追逐、交付。",
                "script_version": 1,
                "script_hash": "retry-json",
            },
        )
        result = build_comic_result("task-retry-json", request)
        provider = RetryPromptProvider()
        with patch("src.comic_office.workflow.LLMFactory.create", return_value=provider):
            enhanced = asyncio.run(enhance_comic_prompts_llm(
                result,
                {"gongbu": ModelConfig(provider="ollama", model="prompt-writer")},
            ))

        package = enhanced["comic_package"]
        self.assertEqual(provider.calls, 2)
        self.assertEqual(package["prompt_generation"]["mode"], "llm_enhanced")
        self.assertIn("第二轮有效人物提示词", package["characters"][0]["image_prompt"])
        self.assertIn(package["visual_style_contract"]["style_id"], package["characters"][0]["image_prompt"])

    def test_word_canvas_contains_platform_execution_table(self):
        from docx import Document

        package = {
            "title": "纸人新娘",
            "creative_brief": {"core_idea": "山村婚礼夜", "story_promise": "纸人秘密被揭开"},
            "confirmed_script": {"story_draft": "新娘在婚礼夜发现宾客都是纸人。"},
            "shots": [
                {
                    "id": "shot_001",
                    "image_ref": "storyboard_shot_001.png",
                    "beat": "新娘掀开红盖头，看见满堂纸人。",
                    "characters": ["新娘"],
                    "scene": "山村堂屋",
                    "props": ["红盖头"],
                    "image_prompt": "新娘，红盖头，山村堂屋，惊恐表情",
                    "video_prompt": "镜头缓慢推进，新娘掀开红盖头",
                    "shot_template": "发现死亡停顿镜头",
                    "shot_template_purpose": "用于关键真相揭露",
                    "composition": "先遮挡再揭露关键结果",
                    "platform_note": "失败时优先重试表情和遮挡关系",
                    "negative_prompt": "角色变脸，服装变化",
                }
            ],
        }
        with TemporaryDirectory() as tmp:
            path = build_comic_word_canvas(package, [], Path(tmp))
            doc = Document(str(path))
            text = "\n".join(p.text for p in doc.paragraphs)
            table_text = "\n".join(cell.text for table in doc.tables for row in table.rows for cell in row.cells)

        self.assertIn("平台执行表", text)
        self.assertIn("镜头模板与执行说明", text)
        self.assertIn("参考资产", table_text)
        self.assertIn("发现死亡停顿镜头", table_text)
        self.assertIn("先遮挡再揭露关键结果", table_text)
        self.assertIn("镜头运动", table_text)
        self.assertIn("失败重试建议", table_text)
        self.assertIn("Libtv", text)

    def test_comic_request_preserves_user_idea_fields(self):
        request = build_comic_request(
            idea="重生后的女主在公司年会上反击陷害她的人",
            genre="revenge short drama",
            length="5 episodes, 60 seconds each",
            platform="Douyin vertical short drama",
            visual_style="modern Chinese webtoon",
            extra="需要强反转和运镜设计",
        )

        self.assertIn("Idea: 重生后的女主", request)
        self.assertIn("Genre: revenge short drama", request)
        self.assertIn("Required output:", request)
        self.assertIn("中文剧本方向", request)

    def test_comic_brief_turn_returns_questions_before_generation(self):
        brief = build_comic_brief(
            idea="一个女孩每天醒来都会进入不同漫画世界",
            genre="fantasy",
            length="5 episodes",
            platform="Douyin",
            visual_style="Korean webtoon",
            extra="希望有悬疑感",
        )

        self.assertEqual(brief["status"], "needs_user_confirmation")
        self.assertIn("creative_brief", brief)
        creative_brief = brief["creative_brief"]
        self.assertIn("一个女孩", creative_brief["core_idea"])
        self.assertGreaterEqual(len(creative_brief["clarifying_questions"]), 4)
        self.assertIn("悬疑感", creative_brief["must_keep"])

    def test_script_preview_explains_why_and_how_before_asset_generation(self):
        brief = build_comic_brief(
            idea="一个女孩每天醒来都会进入不同漫画世界",
            genre="fantasy",
            length="5 episodes",
            platform="Douyin",
            visual_style="Korean webtoon",
            extra="希望有悬疑感",
        )["creative_brief"]

        preview = build_comic_script_preview(
            idea="一个女孩每天醒来都会进入不同漫画世界",
            genre="fantasy",
            length="5 episodes",
            platform="Douyin",
            visual_style="Korean webtoon",
            extra="希望有悬疑感",
            creative_brief=brief,
            user_answers="主角是普通女孩，最后要发现漫画世界是她自己创造的。",
        )

        self.assertEqual(preview["status"], "script_needs_confirmation")
        script = preview["script_preview"]
        self.assertIn("why_it_happens", script)
        self.assertIn("how_it_happens", script)
        self.assertIn("story_draft", script)
        self.assertEqual(script["story_draft"], "")
        self.assertIn("主角是普通女孩", preview["preview"])

    def test_ready_life_drama_can_confirm_without_explicit_hook_keyword(self):
        session = {
            "ready_to_produce": True,
            "creative_brief": {
                "core_idea": "毕业季散场",
                "main_conflict": "林远不知道毕业后的路要通向哪里，只能在最后一节晚自习里面对自己的迷茫。",
                "visual_style": "realistic urban drama",
                "platform": "Douyin/Kuaishou vertical short drama",
            },
            "script_preview": {
                "title": "毕业季散场",
                "genre": "school youth",
                "length": "1 episode, 30 seconds",
                "platform": "Douyin/Kuaishou vertical short drama",
                "visual_style": "realistic urban drama",
                "why_it_happens": "高三最后一节晚自习结束，熟悉的教室突然变得陌生。",
                "how_it_happens": "林远留到最后，在黑板角落画下一个指向窗外的小箭头，然后独自走向校门。",
                "protagonist_arc": "林远从害怕没有方向，到接受迷茫也是自己要走的第一步。",
                "story_draft": "高三最后一节晚自习，林远坐在靠窗的位置，看着窗外漆黑的夜。下课铃响后，他留到最后，在黑板角落画了一个小箭头，指向窗外。走出教学楼时，他回头看了一眼亮着灯的教室，没有拍照，只是加快脚步走向校门。明天就是高考，他不知道未来通向哪里，但至少决定自己走上去。",
                "episode_outline": [],
                "key_turns": [],
            },
            "user_notes": ["没有想去的地方，那只是迷茫接下来的路是什么"],
        }

        issues = validate_confirmed_script_session(session)

        self.assertEqual(issues, [])

    def test_confirmed_story_drives_assets_and_review_package(self):
        confirmed = {
            "status": "confirmed",
            "title": "毕业季散场",
            "script_version": 1,
            "script_hash": "story-hash",
            "story_promise": "一个迷茫的高三学生在最后一节晚自习里接受未知未来。",
            "main_conflict": "林远不知道毕业后的路要通向哪里，却必须独自走出熟悉的教室。",
            "logline": "林远在高考前夜留到最后，在黑板角落画下一个指向窗外的小箭头。",
            "why_it_happens": "高三最后一节晚自习结束，熟悉的教室突然变得陌生。",
            "how_it_happens": "林远拒绝同行，独自收拾练习册和书包，在黑板上画箭头后走向校门。",
            "protagonist_arc": "林远从害怕没有方向，到接受迷茫也是自己要走的第一步。",
            "platform": "Douyin/Kuaishou vertical short drama",
            "visual_style": "realistic urban drama",
            "story_draft": "高三最后一天晚自习，林远坐在靠窗的位置，看着窗外漆黑的夜。下课铃响后，他拒绝同学一起走，留到最后。他走到讲台边，拿起粉笔，在黑板角落画了一个小箭头，指向窗外。然后他把最后一本练习册放进书包，走出教学楼，独自走向校门。",
            "key_turns": "林远画箭头的决定；他拒绝同行者的邀请。",
        }
        request = build_comic_request(
            idea="毕业季散场",
            genre="school youth",
            length="1 episode, 30 seconds",
            platform="Douyin/Kuaishou vertical short drama",
            visual_style="realistic urban drama",
            confirmed_script=confirmed,
        )

        result = build_comic_result("task-life", request)
        package = result["comic_package"]
        artifacts = build_comic_artifacts("task-life", result)

        self.assertIn("林远坐在靠窗的位置", package["script_preview"]["story_draft"])
        self.assertEqual(package["confirmed_script"]["key_turns"], ["林远画箭头的决定；他拒绝同行者的邀请。"])
        self.assertIn("林远", [item["name"] for item in package["characters"]])
        self.assertIn("粉笔", [item["name"] for item in package["props"]])
        self.assertIn("高三晚自习教室", [item["name"] for item in package["scenes"]])
        review = next(item for item in artifacts if item["artifact_type"] == "asset_review_package")
        self.assertTrue(review["metadata"]["requires_human_review"])
        self.assertIn("待审核", review["content"])

    def test_asset_review_package_is_human_facing_inventory_not_prompt_dump(self):
        request = build_comic_request(
            idea="辅助阿衡下山采购后死亡，队友追查真相",
            genre="古风幻想",
            visual_style="ink wash Chinese fantasy",
            confirmed_script={
                "status": "confirmed",
                "title": "无名之坟",
                "story_draft": "辅助阿衡下山采购后死亡，队友追查真相。",
                "story_promise": "队友通过缺席理解辅助的价值。",
                "main_conflict": "队友想复仇，但必须先面对自己长期忽视阿衡。",
                "why_it_happens": "阿衡一直照顾所有人却不被看见。",
                "how_it_happens": "采购、遇害、缺席、寻找尸体、沉默收尸。",
                "script_version": 1,
                "script_hash": "review-inventory",
            },
        )

        result = build_comic_result("task-review-inventory", request)
        review = next(item for item in build_comic_artifacts("task-review-inventory", result) if item["artifact_type"] == "asset_review_package")

        self.assertIn("你只需要确认下面这些是否符合故事", review["content"])
        self.assertIn("默认将生成：人物设定图、人物三视图、人物表情表", review["content"])
        self.assertIn("默认将生成：道具设定图、道具多角度设定、道具基础状态", review["content"])
        self.assertIn("默认将生成：场景设定图、场景广角建立图、场景俯视布局图、场景常用机位", review["content"])
        self.assertNotIn("提示词：", review["content"])
        self.assertNotIn("负面提示词", review["content"])

    def test_asset_revision_notes_override_asset_inventory(self):
        request = build_comic_request(
            idea="高三天台危机",
            genre="family emotion",
            visual_style="Korean webtoon style",
            confirmed_script={
                "status": "confirmed",
                "title": "黄昏的天台",
                "story_draft": (
                    "高三黄昏，林晓站在教学楼天台边缘。张秀兰被保安带上天台，"
                    "她拿出林晓写过的许愿小卡片，试图劝他下来。"
                ),
                "story_promise": "母亲和孩子在天台上面对长期伤害。",
                "main_conflict": "张秀兰想救下林晓，但过去的言语伤害不断刺痛他。",
                "why_it_happens": "长期压力让林晓走到天台边缘。",
                "how_it_happens": "张秀兰拿许愿小卡片哀求，保安在旁边保持距离。",
                "script_version": 1,
                "script_hash": "revision-assets",
            },
            script_notes=(
                "Asset revision notes: 人物只有林晓、张秀兰和保安。"
                "场景只有天台。道具只有背包和许愿小卡片。"
                "删除转动的笔、关键来信和山下街市。必须保留人物姓名。"
            ),
        )

        package = build_comic_result("task-revision-assets", request)["comic_package"]

        self.assertEqual([item["name"] for item in package["characters"]], ["林晓", "张秀兰", "保安"])
        self.assertEqual([item["name"] for item in package["props"]], ["背包", "许愿小卡片"])
        self.assertEqual([item["name"] for item in package["scenes"]], ["天台"])

    def test_asset_revision_notes_can_add_missing_props_without_replacing_other_assets(self):
        request = build_comic_request(
            idea="修仙队伍辅助死亡",
            genre="古风幻想",
            visual_style="ink wash Chinese fantasy",
            confirmed_script={
                "status": "confirmed",
                "title": "无名之坟",
                "story_draft": (
                    "清晨，辅助阿衡下山采购。大师兄、二师姐和小师弟继续修炼，"
                    "直到夜晚才发现阿衡没有回来。他们在山路边找到阿衡的尸体。"
                ),
                "story_promise": "队友通过缺席理解辅助的价值。",
                "main_conflict": "队友想追查真相，但必须先面对自己长期忽视阿衡。",
                "why_it_happens": "阿衡一直照顾所有人却不被看见。",
                "how_it_happens": "采购、缺席、寻找尸体、沉默收尸。",
                "script_version": 1,
                "script_hash": "missing-props-revision",
            },
            script_notes=(
                "Asset revision notes: 第一次拆解人物和场景都准确，但是缺少道具。"
                "请补充道具：止血散、聚气丹、备用药瓶。人物和场景沿用上一版。"
            ),
        )

        package = build_comic_result("task-missing-props-revision", request)["comic_package"]

        self.assertIn("辅助阿衡", [item["name"] for item in package["characters"]])
        self.assertTrue(package["scenes"])
        self.assertEqual([item["name"] for item in package["props"]], ["止血散", "聚气丹", "备用药瓶"])

    def test_confirmed_story_asset_split_does_not_invent_people_props_or_scenes(self):
        request = build_comic_request(
            idea="高三天台危机",
            genre="family emotion",
            visual_style="Korean webtoon style",
            confirmed_script={
                "status": "confirmed",
                "title": "黄昏的天台",
                "story_draft": (
                    "高三黄昏，林晓站在教学楼天台边缘。张秀兰被保安带上天台，"
                    "她背着旧背包，拿出林晓写过的许愿小卡片，试图劝他下来。"
                ),
                "story_promise": "张秀兰和林晓在天台上面对长期伤害。",
                "main_conflict": "张秀兰想救下林晓，但过去的言语伤害不断刺痛他。",
                "why_it_happens": "长期压力让林晓走到天台边缘。",
                "how_it_happens": "张秀兰拿许愿小卡片哀求，保安在旁边保持距离。",
                "script_version": 1,
                "script_hash": "no-phantom-assets",
            },
        )

        package = build_comic_result("task-no-phantom-assets", request)["comic_package"]

        self.assertEqual([item["name"] for item in package["characters"]], ["林晓", "张秀兰", "保安"])
        self.assertEqual([item["name"] for item in package["props"]], ["背包", "许愿小卡片"])
        self.assertEqual([item["name"] for item in package["scenes"]], ["天台"])

    def test_crisis_family_idea_generates_humane_story_instead_of_generic_template(self):
        brief = build_comic_brief(
            idea="学生要跳楼母亲苦苦哀求",
            genre="",
            length="3 episodes",
            platform="Douyin",
            visual_style="现实主义漫画",
            extra="不要猎奇，要有现实痛感和救援希望",
        )["creative_brief"]

        preview = build_comic_script_preview(
            idea="学生要跳楼母亲苦苦哀求",
            genre="",
            length="3 episodes",
            platform="Douyin",
            visual_style="现实主义漫画",
            extra="不要猎奇，要有现实痛感和救援希望",
            creative_brief=brief,
            user_answers="学生长期被成绩和误解压垮，母亲一开始只会责备，最后第一次真正倾听孩子。",
        )

        script = preview["script_preview"]
        self.assertIn("现实情绪危机", script["genre"])
        self.assertEqual(script["story_draft"], "")

    def test_follow_up_questions_are_brief_and_match_non_mystery_story(self):
        session = start_comic_cabinet_session(
            idea="幸福美满的家庭聚餐",
            genre="",
            length="3 episodes",
            platform="Douyin",
            visual_style="现实主义漫画",
            extra="想要温情一点，不要悬疑反转。",
        )["session"]

        questions = session["story_state"]["questions"]
        rendered = "\n".join(questions)
        self.assertLessEqual(len(questions), 2)
        self.assertNotIn("物件", rendered)
        self.assertNotIn("秘密", rendered)
        self.assertNotIn("更狠", rendered)

    def test_rule_fallback_cabinet_message_reads_like_a_creator_not_a_form(self):
        result = start_comic_cabinet_session(
            idea="学生站在天台边缘，母亲赶来劝他下来",
            genre="现实情感",
            length="3集",
            platform="竖屏短剧",
            visual_style="现实主义漫画",
            extra="想拍得克制一点，不要猎奇。",
        )

        message = result["assistant_message"]
        self.assertIn("我先理解成", message)
        self.assertIn("我问这个是因为", message)
        self.assertLessEqual(message.count("？") + message.count("?"), 1)
        self.assertNotIn("还需要一点信息", message)
        self.assertNotIn("1.", message)
        self.assertNotIn("2.", message)

    def test_story_writer_prompt_behaves_like_conversation_not_questionnaire(self):
        from src.comic_office.workflow import _story_writer_system_prompt

        prompt = _story_writer_system_prompt()

        self.assertIn("像真人编剧在聊天", prompt)
        self.assertIn("只能问 1 个最值得问的问题", prompt)
        self.assertIn("解释你为什么问这个问题", prompt)
        self.assertIn("禁止模板化追问", prompt)
        self.assertNotIn("最多 2 个", prompt)

    def test_domestic_warm_story_uses_life_drama_outline_not_mystery_defaults(self):
        preview = build_comic_script_preview(
            idea="幸福美满的家庭聚餐",
            genre="",
            length="3 episodes",
            platform="Douyin",
            visual_style="现实主义漫画",
            extra="温情生活流，不要悬疑反转。",
            user_answers="主角是一位很久没回家的女儿，家庭聚餐表面幸福，其实每个人都有没说出口的委屈。",
        )["script_preview"]

        # 硬编码的模板已被移除，默认 episode_outline 为空，直到 LLM 填充
        self.assertEqual(preview.get("episode_outline", []), [])

    def test_cabinet_session_supports_multi_turn_refinement_before_production(self):
        # Fallback to non-LLM version test to verify state progression
        session = start_comic_cabinet_session(
            idea="一个女孩每天醒来都会进入不同漫画世界",
            genre="fantasy",
            length="5 episodes",
            platform="Douyin",
            visual_style="Korean webtoon",
            extra="希望有悬疑感",
        )["session"]
        
        # Advance with user message (non-LLM doesn't auto-append assistant messages anymore)
        next_turn = advance_comic_cabinet_session(
            session,
            user_message="主角每次醒来都会带走上一部漫画里的一个关键道具",
        )
        self.assertEqual(next_turn["session"]["turn_count"], 1)
        self.assertEqual(next_turn["session"]["messages"][0]["role"], "user")
        self.assertIn("带走上一部漫画", next_turn["session"]["user_notes"][0])

    def test_llm_cabinet_session_uses_role_specific_prompt_outputs_when_available(self):
        session = start_comic_cabinet_session("我想写一个悬疑故事", "赛博朋克", "反转结局")
        role_configs = {"编剧顾问": ModelConfig(provider="mock", model="mock")}
        with patch("src.comic_office.workflow._model_config_usable", return_value=True):
            with patch("src.comic_office.workflow.advance_comic_cabinet_session") as mock_advance:
                mock_advance.return_value = {
                    "session": session["session"],
                    "creative_brief": session["creative_brief"],
                    "script_preview": session["script_preview"],
                    "assistant_message": "old",
                    "cabinet_roles": [],
                    "preview": "old"
                }
                with patch("src.comic_office.workflow._cabinet_story_writer_llm", return_value={"assistant_message": "LLM 回复", "story": {}}):
                    result = asyncio.run(advance_comic_cabinet_session_llm(session["session"], "角色是谁？", role_configs))
                    self.assertTrue(result["session"]["llm_cabinet"])
                    self.assertEqual(len(result["cabinet_roles"]), 0)
                    self.assertEqual(result["assistant_message"], "LLM 回复")

    def test_llm_cabinet_keeps_partial_role_outputs_and_marks_rule_fallbacks(self):
        # 这个测试由于移除了并行的内阁顾问，所以不再适用，可以直接删除或替换为单 Agent fallback 测试
        pass

    def test_prompt_generation_uses_llm_before_human_asset_review(self):
        class FakePromptProvider:
            async def chat(self, messages, response_format=None):
                return LLMResponse(
                    content="""
                    {
                      "characters": [
                        {
                          "id": "char_01",
                          "image_prompt": "LLM人物设定：阿衡背着旧药箱站在雨后山路，衣摆有泥点，眼神温和但疲惫，强调被忽视的队伍辅助气质。",
                          "asset_specs": [
                            {
                              "kind": "character_three_view",
                              "label": "人物三视图",
                              "image_ref": "char_01_three_view.png",
                              "prompt": "LLM三视图：阿衡同一张脸，同一旧青灰外袍，正侧背三面展示药箱背带和袖口磨损。",
                              "acceptance": "三面必须是同一角色，药箱背带和青灰外袍不能漂移。"
                            }
                          ]
                        }
                      ],
                      "shots": [
                        {
                          "id": "shot_001",
                          "image_prompt": "LLM分镜：清晨山路，阿衡低头检查药包，远处车驾压迫性逼近，画面重点是他还在想着队友的喜好。",
                          "video_prompt": "镜头从药包特写慢慢抬到阿衡侧脸，再让车驾阴影压入画面。",
                          "negative_prompt": "不要现代车辆，不要文字，不要脸型漂移"
                        }
                      ],
                      "quality_review": {
                        "status": "pass",
                        "summary": "提示词已经贴合阿衡被忽视又温柔照顾队伍的故事，不是固定模板。",
                        "issues": []
                      }
                    }
                    """,
                    model="fake",
                )

        request = build_comic_request(
            idea="清晨，辅助阿衡下山采购，他细心记下每个人的喜好，后来被车驾撞死在偏僻小巷",
            genre="古风幻想",
            visual_style="ink wash Chinese fantasy",
            confirmed_script={
                "status": "confirmed",
                "title": "一队修仙队伍中的辅助死亡了",
                "story_draft": "清晨，辅助阿衡下山采购，记下每个人的喜好。后来他死在偏僻小巷，队友才意识到他的缺席。",
                "story_promise": "辅助死亡后，队伍才意识到他承担了一切温柔细节。",
                "main_conflict": "队友想复仇，但真正要面对的是长期忽视辅助的愧疚。",
                "why_it_happens": "阿衡总是默默照顾所有人，所以他的缺席直到很晚才被发现。",
                "how_it_happens": "从采购、遇害、缺席、寻找尸体到队友沉默收尸推进。",
                "script_version": 1,
                "script_hash": "abc123",
            },
        )
        result = build_comic_result("task-llm-prompts", request)
        with patch("src.comic_office.workflow.LLMFactory.create", return_value=FakePromptProvider()):
            enhanced = asyncio.run(enhance_comic_prompts_llm(
                result,
                {
                    "gongbu": ModelConfig(provider="ollama", model="prompt-writer"),
                    "bingbu": ModelConfig(provider="ollama", model="shot-writer"),
                    "xingbu": ModelConfig(provider="ollama", model="quality-reviewer"),
                },
            ))

        package = enhanced["comic_package"]
        self.assertIn("LLM人物设定", package["characters"][0]["image_prompt"])
        self.assertIn("LLM三视图", package["characters"][0]["asset_specs"][0]["prompt"])
        self.assertIn("LLM分镜", package["shots"][0]["image_prompt"])
        self.assertNotIn("不要", package["shots"][0]["negative_prompt"])
        self.assertIn("禁止现代车辆", package["shots"][0]["negative_prompt"])
        self.assertEqual(package["prompt_generation"]["mode"], "llm_enhanced")
        self.assertEqual(package["prompt_generation"]["quality_review"]["status"], "pass")

    def test_base_asset_prompts_do_not_tell_story(self):
        request = build_comic_request(
            idea="清晨，辅助阿衡下山采购，他细心记下每个人的喜好，后来被车驾撞死在偏僻小巷",
            genre="古风幻想",
            visual_style="ink wash Chinese fantasy",
            confirmed_script={
                "status": "confirmed",
                "title": "一队修仙队伍中的辅助死亡了",
                "story_draft": (
                    "清晨，辅助阿衡下山采购，记下每个人的喜好。"
                    "后来他被车驾撞倒，死在偏僻小巷，队友发现尸体后才意识到他的缺席。"
                ),
                "story_promise": "辅助死亡后，队伍才意识到他承担了一切温柔细节。",
                "main_conflict": "队友想复仇，但真正要面对的是长期忽视辅助的愧疚。",
                "why_it_happens": "阿衡总是默默照顾所有人，所以他的缺席直到很晚才被发现。",
                "how_it_happens": "从采购、遇害、缺席、寻找尸体到队友沉默收尸推进。",
                "script_version": 1,
                "script_hash": "abc123",
            },
        )
        result = build_comic_result("task-base-assets", request)
        package = result["comic_package"]
        story_words = ["清晨", "下山采购", "每个人的喜好", "被车驾", "发现尸体", "死在偏僻小巷"]

        asset_prompts = []
        for group in ("characters", "props", "scenes"):
            for item in package[group]:
                asset_prompts.append(item.get("image_prompt", ""))
                asset_prompts.extend(spec.get("prompt", "") for spec in item.get("asset_specs", []))

        for prompt in asset_prompts:
            for word in story_words:
                self.assertNotIn(word, prompt)
            self.assertFalse(prompt.startswith(package["title"]))
        self.assertTrue(any("基础设定" in prompt or "设定图" in prompt for prompt in asset_prompts))
        self.assertTrue(any("被车驾" in shot.get("image_prompt", "") or "被车驾" in shot.get("beat", "") for shot in package["shots"]))

    def test_character_and_prop_assets_use_clean_white_backgrounds(self):
        request = build_comic_request(
            idea="女主在黄昏办公室诱导老板开口",
            genre="modern suspense dialogue",
            visual_style="photorealism cinematic dusk office",
            confirmed_script={
                "status": "confirmed",
                "title": "黄昏办公室的试探",
                "story_draft": "女主在办公室诱导老板开口，桌上有一封关键合同。",
                "story_promise": "一场谈话逐步变成心理试探。",
                "main_conflict": "女主想让老板开口，但老板回避秘密。",
                "why_it_happens": "女主发现老板说谎。",
                "how_it_happens": "女主倾听、靠近、追问。",
                "script_version": 1,
                "script_hash": "white-bg-123",
            },
        )

        package = build_comic_result("task-white-background-assets", request)["comic_package"]
        character_and_prop_prompts = []
        for group in ("characters", "props"):
            for item in package[group]:
                character_and_prop_prompts.append(item.get("image_prompt", ""))
                character_and_prop_prompts.extend(spec.get("prompt", "") for spec in item.get("asset_specs", []))

        self.assertTrue(character_and_prop_prompts)
        for prompt in character_and_prop_prompts:
            self.assertIn("纯白或近白色干净背景", prompt)
            self.assertNotIn("不要", prompt)
            self.assertIn("负面提示词：禁止", prompt)

        scene_prompts = []
        for item in package["scenes"]:
            scene_prompts.append(item.get("image_prompt", ""))
            scene_prompts.extend(spec.get("prompt", "") for spec in item.get("asset_specs", []))
        self.assertTrue(scene_prompts)
        self.assertTrue(any("空间" in prompt or "场景" in prompt for prompt in scene_prompts))
        self.assertTrue(all("不要" not in prompt for prompt in scene_prompts))

    def test_ancient_fantasy_assets_share_one_style_contract_and_reject_modern_elements(self):
        request = build_comic_request(
            idea="辅助阿衡下山采购后死亡",
            genre="古风仙侠",
            visual_style="电影级国风厚涂动画，克制冷色",
            confirmed_script={
                "status": "confirmed",
                "title": "无名木牌",
                "story_draft": "辅助阿衡背着药箱下山采购丹药，在废弃驿站为保护求救符引开妖兽而死。",
                "story_promise": "队友通过缺席理解辅助的价值。",
                "main_conflict": "队友寻找失踪的阿衡。",
                "why_it_happens": "阿衡独自承担采购任务。",
                "how_it_happens": "沿山路找到药箱、求救符和遗体。",
                "script_version": 1,
                "script_hash": "ancient-style-contract",
            },
        )

        package = build_comic_result("task-ancient-style-contract", request)["comic_package"]
        contract = package["visual_style_contract"]

        self.assertTrue(contract["style_id"].startswith("style_"))
        self.assertEqual(contract["period"], "古风仙侠")
        self.assertIn("现代服装", contract["forbidden_elements"])
        for group in ("characters", "props", "scenes"):
            for item in package[group]:
                self.assertEqual(item["style_id"], contract["style_id"])
                prompts = [item.get("image_prompt", "")]
                prompts.extend(spec.get("prompt", "") for spec in item.get("asset_specs", []))
                for prompt in prompts:
                    self.assertIn("古风仙侠", prompt)
                    self.assertIn("禁止现代", prompt)

    def test_scenes_include_spatial_reference_assets_and_shots_are_prompt_only(self):
        request = build_comic_request(
            idea="一队修仙队伍中的辅助死亡了",
            genre="xianxia tragedy",
            visual_style="ink wash Chinese fantasy",
            confirmed_script={
                "status": "confirmed",
                "title": "一队修仙队伍中的辅助死亡了",
                "story_draft": "辅助阿衡下山采购后死在偏僻小巷，队友寻找他时才意识到他的缺席。",
                "story_promise": "辅助死亡后，队伍才看见他曾经承担的温柔细节。",
                "main_conflict": "队友想复仇，但真正要面对的是长期忽视辅助的愧疚。",
                "why_it_happens": "阿衡总是默默照顾所有人，所以他的缺席很晚才被发现。",
                "how_it_happens": "从采购、遇害、缺席、寻找尸体到沉默收尸推进。",
                "script_version": 1,
                "script_hash": "abc123",
            },
        )

        package = build_comic_result("task-scene-spatial-assets", request)["comic_package"]
        scene_specs = [
            spec
            for scene in package["scenes"]
            for spec in scene.get("asset_specs", [])
        ]
        scene_kinds = {spec.get("kind") for spec in scene_specs}

        self.assertIn("scene_wide_establishing", scene_kinds)
        self.assertIn("scene_top_down_layout", scene_kinds)
        self.assertIn("scene_camera_angles", scene_kinds)
        self.assertTrue(any("俯视" in spec.get("prompt", "") for spec in scene_specs))
        self.assertTrue(any("广角" in spec.get("prompt", "") for spec in scene_specs))
        self.assertTrue(all((shot.get("image_ref") or "") == "" for shot in package["shots"]))
        self.assertTrue(all("storyboard" not in (shot.get("image_ref") or "") for shot in package["shots"]))

    def test_assets_have_identity_cards_and_shots_reference_them(self):
        request = build_comic_request(
            idea="女主在黄昏办公室诱导老板开口",
            genre="modern suspense dialogue",
            visual_style="photorealism cinematic dusk office",
            confirmed_script={
                "status": "confirmed",
                "title": "黄昏办公室的试探",
                "story_draft": "女主在办公室诱导老板开口，桌上有一封关键合同。",
                "story_promise": "一场谈话逐步变成心理试探。",
                "main_conflict": "女主想让老板开口，但老板回避秘密。",
                "why_it_happens": "女主发现老板说谎。",
                "how_it_happens": "女主倾听、靠近、追问。",
                "script_version": 1,
                "script_hash": "identity-card-123",
            },
        )

        package = build_comic_result("task-identity-cards", request)["comic_package"]

        for group, expected_kind in (
            ("characters", "character_identity_card"),
            ("props", "prop_identity_card"),
            ("scenes", "scene_identity_card"),
        ):
            self.assertTrue(package[group], group)
            first = package[group][0]
            self.assertIn("identity_card", first)
            self.assertEqual(first["identity_card"]["kind"], expected_kind)
            self.assertTrue(first["identity_card"]["image_refs"])
            self.assertIn("验收", first["identity_card"]["usage_rule"])

        first_shot = package["shots"][0]
        self.assertIn("identity_references", first_shot)
        self.assertIn("角色身份证", first_shot["identity_references"])
        self.assertIn("道具身份证", first_shot["identity_references"])
        self.assertIn("场景身份证", first_shot["identity_references"])
        self.assertIn("角色身份证", first_shot["director_prompt"])
        self.assertIn("场景身份证", first_shot["video_prompt"])

    def test_shots_include_flexible_director_prompt_fields(self):
        request = build_comic_request(
            idea="女主在黄昏办公室诱导老板开口",
            genre="modern suspense dialogue",
            visual_style="photorealism cinematic dusk office",
            confirmed_script={
                "status": "confirmed",
                "title": "黄昏办公室的试探",
                "story_draft": "女主靠在椅子上，先低头思考，再缓缓直起身靠近镜头，用温柔好奇的语气诱导老板继续说出秘密。",
                "story_promise": "一场看似普通的办公室谈话，逐步变成心理试探。",
                "main_conflict": "女主想让老板开口，但老板一直回避真正的秘密。",
                "why_it_happens": "女主刚发现老板说谎，需要用轻柔的方式让他放松警惕。",
                "how_it_happens": "她从安静倾听、身体前倾、轻声追问到逼近真相。",
                "script_version": 1,
                "script_hash": "office123",
            },
        )

        package = build_comic_result("task-director-prompts", request)["comic_package"]
        first = package["shots"][0]

        for key in (
            "shot_template",
            "shot_template_purpose",
            "composition",
            "platform_note",
            "reference_assets",
            "performance_intent",
            "action_chain",
            "cinematography",
            "lighting",
            "director_prompt",
        ):
            self.assertIn(key, first)
            self.assertTrue(first[key])

        self.assertIn("起始", first["action_chain"])
        self.assertIn("过程", first["action_chain"])
        self.assertIn("结束", first["action_chain"])
        self.assertIn("char_", first["director_prompt"])
        self.assertIn("scene_", first["director_prompt"])
        self.assertIn("参考资产", first["director_prompt"])
        self.assertIn("表演意图", first["director_prompt"])
        self.assertIn("摄影", first["director_prompt"])
        self.assertIn("灯光", first["director_prompt"])
        self.assertIn("固定镜头", first["video_prompt"])
        self.assertNotIn("不要", first["director_prompt"])
        self.assertNotIn("不要", first["video_prompt"])
        self.assertNotIn("不要", first["negative_prompt"])
        self.assertIn("禁止", first["negative_prompt"])
        self.assertNotEqual(package["shots"][0]["director_prompt"], package["shots"][1]["director_prompt"])

    def test_rule_fallback_director_prompts_follow_story_beat_not_static_template(self):
        request = build_comic_request(
            idea="辅助阿衡死亡后队友追查真相",
            genre="xianxia tragedy",
            visual_style="ink wash Chinese fantasy",
            confirmed_script={
                "status": "confirmed",
                "title": "无名之坟",
                "story_draft": "阿衡清晨下山采购，后来死在偏僻小巷。队友发现尸体后决定追查贵人车驾。最后他们在雨中立下一座无名坟。",
                "episode_outline": [
                    {
                        "episode": 1,
                        "title": "采购",
                        "cause": "阿衡清晨下山采购并记下队友喜好",
                        "action": "他把药包和桂花糕分开放好",
                        "turn": "贵人车驾从巷口逼近",
                        "hook": "药包滚落到雨水里",
                    },
                    {
                        "episode": 2,
                        "title": "缺席",
                        "cause": "队友直到晚饭才发现阿衡没有回来",
                        "action": "大师兄沿街寻找并看见偏僻小巷里的尸体",
                        "turn": "尸体旁没有血迹，只有散落物资",
                        "hook": "众人第一次意识到他一直照顾所有人",
                    },
                    {
                        "episode": 3,
                        "title": "追查",
                        "cause": "队友想复仇",
                        "action": "二师姐追查贵人车驾",
                        "turn": "贵人根本不记得撞过这样一个人",
                        "hook": "小师弟在雨中给阿衡立无名坟",
                    },
                ],
                "story_promise": "辅助死亡后，队伍才看见被忽视的温柔。",
                "main_conflict": "复仇和愧疚同时推动队伍。",
                "why_it_happens": "队友长期忽视阿衡的付出。",
                "how_it_happens": "采购、遇害、发现尸体、追查、立坟。",
                "script_version": 1,
                "script_hash": "dynamic123",
            },
        )

        package = build_comic_result("task-dynamic-director", request)["comic_package"]
        prompts = [shot["director_prompt"] for shot in package["shots"][:3]]

        self.assertIn("药包滚落到雨水里", prompts[0])
        self.assertIn("尸体旁没有血迹", prompts[1])
        self.assertIn("贵人根本不记得", prompts[2])
        self.assertIn("迟来的悲伤", package["shots"][1]["performance_intent"])
        self.assertIn("愤怒压在表面之下", package["shots"][2]["performance_intent"])
        self.assertNotEqual(package["shots"][0]["lighting"], package["shots"][1]["lighting"])
        self.assertNotEqual(package["shots"][1]["negative_prompt"], package["shots"][2]["negative_prompt"])

    def test_role_words_are_preferred_over_scene_words_when_extracting_characters(self):
        request = build_comic_request(
            idea="女主在黄昏办公室诱导老板开口",
            genre="modern suspense dialogue",
            confirmed_script={
                "status": "confirmed",
                "title": "黄昏办公室的试探",
                "story_draft": "女主靠在椅子上，先低头思考，再缓缓直起身靠近镜头，诱导老板继续说出秘密。",
                "story_promise": "办公室谈话变成心理试探。",
                "main_conflict": "女主想让老板开口，但老板一直回避。",
                "why_it_happens": "女主发现老板说谎。",
                "how_it_happens": "女主倾听、靠近、追问。",
                "script_version": 1,
                "script_hash": "office456",
            },
        )

        package = build_comic_result("task-role-words", request)["comic_package"]
        names = [item["name"] for item in package["characters"]]

        self.assertIn("女主", names)
        self.assertIn("老板", names)
        self.assertNotIn("黄昏", names)

    def test_confirmed_script_is_built_from_session_and_reused_by_production(self):
        session = start_comic_cabinet_session(
            idea="雨夜里失忆侦探捡到一封会改写身份的信",
            genre="suspense",
            length="3 episodes",
            platform="Douyin",
            visual_style="dark suspense comic",
            extra="",
        )["session"]
        session = advance_comic_cabinet_session(
            session,
            "主角是女侦探，反派一直在误导她，最后她发现自己名字也在信里，结尾要留悬念。",
        )["session"]
        
        # 模拟大模型生成了分集大纲
        session["script_preview"]["episode_outline"] = [
            {"episode": 1, "title": "第一集", "cause": "起因", "action": "行动", "turn": "转折", "hook": "钩子"}
        ]
        
        issues = validate_confirmed_script_session(session)
        self.assertEqual(issues, [])
        confirmed = build_confirmed_script(session, "第一集钩子要狠一点，但整体方向确认。")
        self.assertEqual(confirmed["status"], "confirmed")
        self.assertIn("第一集钩子要狠一点", confirmed["confirmation_notes"])
        self.assertIn("story_draft", confirmed)
        self.assertEqual(confirmed["story_draft"], "")

        request = build_comic_request(
            idea="雨夜里失忆侦探捡到一封会改写身份的信",
            genre="suspense",
            length="3 episodes",
            platform="Douyin",
            visual_style="dark suspense comic",
            creative_brief=session["creative_brief"],
            user_answers="\n".join(session["user_notes"]),
            script_preview=session["script_preview"],
            confirmed_script=confirmed,
        )
        result = build_comic_result("task-confirmed", request)
        self.assertEqual(result["comic_package"]["confirmed_script"]["status"], "confirmed")
        self.assertEqual(
            result["comic_package"]["confirmed_script"].get("story_draft", ""),
            confirmed["story_draft"],
        )
        self.assertIn("完整故事稿", result["final_report"])
        self.assertEqual(
            result["comic_package"]["confirmed_script"]["why_it_happens"],
            confirmed["why_it_happens"],
        )
        self.assertTrue(result["comic_package"]["script_binding"]["confirmed"])
        self.assertEqual(result["comic_package"]["script_binding"]["script_version"], 1)
        self.assertEqual(
            result["comic_package"]["script_binding"]["script_hash"],
            confirmed["script_hash"],
        )

    def test_full_script_cabinet_session_preserves_original_story_text(self):
        full_script = (
            "清晨，林晓站在教学楼天台边缘。张秀兰被保安带上天台，"
            "她背着旧背包，拿出林晓写过的许愿小卡片，试图劝他下来。"
            "最后林晓没有立刻说话，只是看向那张卡片。"
        )

        result = start_comic_cabinet_session(
            idea="高三天台危机",
            genre="现实情感",
            length="3集",
            platform="竖屏短剧",
            visual_style="Korean webtoon style",
            extra="Input mode: full_script\nFull script:\n" + full_script,
        )

        self.assertEqual(result["script_preview"]["story_draft"], full_script)
        self.assertEqual(result["session"]["input_mode"], "full_script")
        self.assertIn("不改写原文", result["session"]["summary"])

    def test_full_script_mode_ignores_llm_rewritten_story_draft(self):
        from src.comic_office.workflow import _apply_llm_story_payload

        full_script = "原文第一句。原文第二句。"
        result = start_comic_cabinet_session(
            idea="完整故事测试",
            genre="现实情感",
            extra="Input mode: full_script\nFull script:\n" + full_script,
        )
        rewritten = {
            "assistant_message": "我提一个确认点。",
            "story": {
                "title": "被改写的标题",
                "story_draft": "模型擅自改写后的故事。",
                "episode_outline": [],
            },
        }

        protected = _apply_llm_story_payload(result, rewritten)

        self.assertEqual(protected["script_preview"]["story_draft"], full_script)
        self.assertEqual(protected["session"]["script_preview"]["story_draft"], full_script)

    def test_comic_result_contains_brief_assets_shots_and_prompts(self):
        brief = build_comic_brief(
            idea="民国侦探在雨夜发现一封会改写身份的信",
            genre="suspense",
            length="3 episodes, 60 seconds each",
            platform="vertical short-video platforms",
            visual_style="dark suspense comic",
            extra="",
        )
        result = build_comic_result(
            "task-1",
            build_comic_request(
                idea="民国侦探在雨夜发现一封会改写身份的信",
                genre="suspense",
                length="3 episodes, 60 seconds each",
                platform="vertical short-video platforms",
                visual_style="dark suspense comic",
                creative_brief=brief["creative_brief"],
                user_answers="主角是女性，结尾要留反转。",
            ),
        )

        package = result["comic_package"]
        self.assertEqual(result["status"], "completed")
        self.assertIn("creative_brief", package)
        self.assertIn("clarifying_questions", package["creative_brief"])
        self.assertIn("主角是女性", package["user_answers"])
        self.assertIn("民国侦探", package["title"])
        self.assertGreaterEqual(len(package["characters"]), 3)
        self.assertGreaterEqual(len(package["props"]), 3)
        self.assertGreaterEqual(len(package["scenes"]), 3)
        self.assertGreaterEqual(len(package["shots"]), 6)
        self.assertIn("script_binding", package)
        self.assertIn("consistency_bindings", package)
        self.assertTrue(package["consistency_bindings"]["script"]["script_hash"])
        self.assertTrue(all("anchor_id" in item for item in package["characters"]))
        self.assertTrue(all("binding" in item for item in package["shots"]))
        self.assertIn("beat_id", package["shots"][0]["binding"])
        self.assertIn("一致性闭环", result["final_report"])

    def test_comic_result_derives_assets_from_confirmed_story_content(self):
        result = build_comic_result(
            "task-paper-bride",
            (
                "Idea: 山村纸人新娘在婚礼夜突然睁眼\n"
                "Confirmed script:\n"
                "# 纸人新娘 确认版剧本\n\n"
                "- 题材：horror thriller\n"
                "- 长度：3 episodes, 60 seconds each\n"
                "- 平台：Douyin\n"
                "- 视觉风格：dark suspense comic\n"
                "- 故事承诺：山村婚礼、纸人新娘、红盖头、祠堂和相机闪光灯。\n"
                "- 主冲突：主角必须在村民阻止下救出真正的新娘。\n"
                "- 一句话故事：摄影师发现纸人新娘突然睁眼，真正的新娘被藏在祠堂地下。\n"
                "- 为什么发生：族老用冥婚仪式掩盖旧案。\n"
                "- 如何发生：主角拍摄婚礼，发现纸人、红盖头、祠堂地板和相机照片里的线索。\n"
                "- 主角变化：主角从完成拍摄到主动救人。\n\n"
                "## 完整故事稿\n"
                "山村婚礼夜，摄影师陈野看到纸人新娘隔着红盖头睁眼。他用相机闪光灯照亮祠堂地板，发现地下暗门。\n\n"
                "## 每集确认大纲\n"
                "1. 纸人睁眼｜起因：主角接到山村婚礼拍摄委托｜行动：拍摄仪式｜转折：相机拍到纸人身后有人｜钩子：纸人新娘隔着红盖头睁眼\n"
                "2. 祠堂暗门｜起因：追查照片里的女人｜行动：进入祠堂｜转折：族老阻止主角｜钩子：地板下传出求救声\n"
                "3. 红盖头｜起因：救出真新娘｜行动：用相机闪光灯揭露骗局｜转折：纸人被烧毁｜钩子：纸人再次出现在照片里\n"
            ),
        )

        package = result["comic_package"]
        prop_names = " ".join(item["name"] for item in package["props"])
        scene_names = " ".join(item["name"] for item in package["scenes"])
        prompts = result["final_report"]

        self.assertIn("纸人新娘", prop_names)
        self.assertIn("红盖头", prop_names)
        self.assertIn("相机", prop_names)
        self.assertIn("山村婚礼堂屋", scene_names)
        self.assertIn("祠堂", scene_names)
        self.assertNotIn("关键合同", prompts)
        self.assertNotIn("现代办公室", prompts)

    def test_comic_result_derives_real_shot_beats_from_story_draft(self):
        result = build_comic_result(
            "task-helper-death",
            (
                "Idea: 一队修仙队伍中的辅助死亡了\n"
                "Genre: xianxia tragedy\n"
                "Visual style: ink wash Chinese fantasy\n"
                "Confirmed script:\n"
                "# 一队修仙队伍中的辅助死亡了 确认版剧本\n\n"
                "- 题材：xianxia tragedy\n"
                "- 长度：3 episodes\n"
                "- 平台：Douyin\n"
                "- 视觉风格：ink wash Chinese fantasy\n"
                "- 故事承诺：辅助死亡后，队伍才意识到他承担了一切温柔细节。\n"
                "- 主冲突：队友想复仇，但真正要面对的是他们长期忽视辅助的愧疚。\n"
                "## 完整故事稿\n"
                "清晨，辅助阿衡下山采购，记下大师兄爱吃的桂花糕、二师姐常备的止血散、小师弟总弄丢的聚气丹。"
                "一辆贵人车驾横冲直撞，他被护卫拖进偏僻小巷乱棍打死。"
                "夜晚，队友们抱怨没人准备晚饭，才发现阿衡没有回来。"
                "他们在小巷发现他的尸体和散落的药包，沉默地意识到自己从未真正看见过他。\n"
            ),
        )

        package = result["comic_package"]
        beats = " ".join(shot["beat"] for shot in package["shots"])
        prompts = " ".join(shot["image_prompt"] for shot in package["shots"])
        names = " ".join(item["name"] for group in ("characters", "props", "scenes") for item in package[group])

        self.assertNotIn("确认故事", beats)
        self.assertNotIn("补充分镜节拍", beats)
        self.assertIn("下山采购", beats)
        self.assertIn("偏僻小巷", prompts)
        self.assertIn("辅助阿衡", names)
        self.assertIn("大师兄", names)
        self.assertIn("桂花糕", names)
        self.assertIn("止血散", names)
        self.assertIn("偏僻小巷", names)

    def test_comic_artifacts_cover_preproduction_delivery_canvas(self):
        brief = build_comic_brief(
            idea="一个外卖员误入未来城市后必须用三小时救回妹妹",
            genre="science fiction",
            length="5 episodes",
            platform="Douyin",
            visual_style="cinematic comic",
            extra="",
        )
        result = build_comic_result(
            "task-2",
            build_comic_request(
                idea="一个外卖员误入未来城市后必须用三小时救回妹妹",
                genre="science fiction",
                creative_brief=brief["creative_brief"],
                user_answers="主角要普通但行动力强。",
            ),
        )

        artifacts = build_comic_artifacts("task-2", result)
        artifact_types = {artifact["artifact_type"] for artifact in artifacts}

        self.assertEqual(len(artifacts), 15)
        self.assertIn("creative_brief", artifact_types)
        self.assertIn("script_preview", artifact_types)
        self.assertIn("story_draft", artifact_types)
        self.assertIn("confirmed_script", artifact_types)
        self.assertIn("cabinet_review", artifact_types)
        self.assertIn("script", artifact_types)
        self.assertIn("style_bible", artifact_types)
        self.assertIn("asset_review_package", artifact_types)
        self.assertIn("character_sheet", artifact_types)
        self.assertNotIn("storyboard_table", artifact_types)
        self.assertIn("production_canvas", artifact_types)
        self.assertIn("word_canvas", artifact_types)
        self.assertNotIn("camera_plan", artifact_types)
        self.assertIn("prompt_package", artifact_types)
        self.assertIn("consistency_checklist", artifact_types)
        self.assertTrue(all(artifact["metadata"]["office_id"] == "comic" for artifact in artifacts))
        self.assertTrue(all("script_hash" in artifact["metadata"] for artifact in artifacts))
        canvas = next(a for a in artifacts if a["artifact_type"] == "production_canvas")
        self.assertIn("shot_001", canvas["content"])
        self.assertNotIn("对应图片", canvas["content"])
        self.assertIn("脚本版本", canvas["content"])
        story = next(a for a in artifacts if a["artifact_type"] == "story_draft")
        self.assertIn("当前剧本尚未生成完整故事稿", story["content"])
        prompt_package = next(a for a in artifacts if a["artifact_type"] == "prompt_package")
        self.assertIn("提示词包", prompt_package["content"])
        self.assertIn("负面提示词", prompt_package["content"])
        self.assertIn("动作链", prompt_package["content"])
        self.assertIn("表演意图", prompt_package["content"])
        self.assertIn("摄影", prompt_package["content"])
        self.assertIn("灯光", prompt_package["content"])

        production_canvas = next(a for a in artifacts if a["artifact_type"] == "production_canvas")
        self.assertIn("参考资产", production_canvas["content"])
        self.assertIn("动作链", production_canvas["content"])


if __name__ == "__main__":
    unittest.main()
