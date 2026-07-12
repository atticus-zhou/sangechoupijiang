import base64
import json
import tempfile
import unittest
from pathlib import Path

from docx import Document

from src.comic_office.v2.asset_manifest import build_asset_manifest
from src.comic_office.v2.contracts import build_contract_bundle
from src.comic_office.v2.production import ImageProductionResult, ImageRecord, PromptPackage
from src.comic_office.v2.prompt_director import PromptPlan, ShotCard


PNG_1X1 = base64.b64decode(
    "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mP8/x8AAusB9Y9ZrWQAAAAASUVORK5CYII="
)
STORY = "林昭走进月塔，最终熄灭月塔。"


def parts(image_dir: Path):
    bundle = build_contract_bundle(STORY, {
        "title": "借月人",
        "genre": "古风幻想",
        "theme": "记忆与光明的代价",
        "protagonist_goal": "熄灭月塔",
        "main_conflict": "月塔燃烧记忆维持光明",
        "causal_chain": ["进入月塔", "作出选择", "熄灭月塔"],
        "ending": "林昭最终熄灭月塔",
        "episodes": [{"episode": 1, "summary": "熄灭月塔", "evidence_quote": "林昭走进月塔"}],
        "visual": {
            "medium": "电影级国风厚涂动画",
            "era": "架空古代",
            "aspect_ratio": "9:16",
            "palette": ["靛青", "银白", "暗朱红"],
            "lighting": "冷月光与暖灯火对照",
            "camera_language": "克制稳定",
            "character_rules": ["脸型固定"],
            "costume_rules": ["古代窄袖长袍"],
            "prop_rules": ["材质固定"],
            "architecture_rules": ["木石结构"],
            "visual_motifs": ["裂纹月灯"],
            "prohibited_elements": ["现代车辆"],
        },
    })
    manifest = build_asset_manifest(bundle, [{
        "asset_type": "character",
        "name": "林昭",
        "evidence_quote": "林昭走进月塔",
        "scene_ids": ["scene_01"],
        "story_purpose": "完成最终选择",
        "visual_locks": ["靛青长袍"],
        "allowed_changes": ["表情", "姿势"],
    }])
    asset = manifest.items[0]
    prompts = tuple(
        PromptPlan(
            object_id=asset.asset_id,
            image_kind=kind,
            purpose="identity_reference",
            generator_prompt=f"林昭 {kind}，靛青长袍，纯白干净背景",
            negative_prompt=("禁止文字水印", "禁止现代服装"),
            style_id=bundle.visual.style_id,
        )
        for kind in asset.planned_images
    )
    shot = ShotCard(
        shot_id="shot_01",
        scene_id="scene_01",
        story_beat="林昭走进月塔",
        reference_asset_ids=(asset.asset_id,),
        action_chain=("林昭推开塔门", "她抬头望向塔心"),
        performance_intent="恐惧逐渐转为决绝",
        framing="中近景平视",
        camera_movement="缓慢前推",
        lighting="冷月光压住暖灯火",
        dialogue="林昭：到此为止。",
        sound="风声与灯芯爆裂声",
        generator_prompt="首帧参考人物资产，林昭推开塔门，缓慢前推。",
        negative_prompt=("禁止脸型变化",),
        retry_strategy="优先修正脸型和动作顺序",
        retry_strategy_label="失败重试",
        style_id=bundle.visual.style_id,
        evidence_quote="林昭走进月塔",
        acceptance_criteria=(
            "首帧必须引用已批准人物资产",
            "动作顺序必须保持推门、抬头",
        ),
        platform_note="适合图生视频首帧参考，先绑定参考图片再粘贴视频提示词。",
    )
    package = PromptPackage(
        package_id="prompts_delivery",
        story_id=bundle.creative.story_id,
        story_version=bundle.creative.story_version,
        style_id=bundle.visual.style_id,
        style_version=bundle.visual.style_version,
        manifest_id=manifest.manifest_id,
        manifest_version=manifest.version,
        prompts=prompts,
        shots=(shot,),
    )
    records = []
    for index, kind in enumerate(asset.planned_images):
        path = image_dir / f"{asset.asset_id}_{kind}.png"
        path.write_bytes(PNG_1X1)
        records.append(ImageRecord(
            image_id=f"img_{asset.asset_id}_{kind}",
            asset_id=asset.asset_id,
            image_kind=kind,
            prompt_hash=f"hash-{kind}",
            path=str(path),
            provider="doubao",
            model="seedream",
            attempts=1,
            status="approved",
            is_identity_baseline=index == 0,
            reference_image_ids=() if index == 0 else (f"img_{asset.asset_id}_{asset.planned_images[0]}",),
            story_id=bundle.creative.story_id,
            story_version=bundle.creative.story_version,
            style_id=bundle.visual.style_id,
            style_version=bundle.visual.style_version,
            manifest_version=manifest.version,
            review={"status": "pass"},
        ))
    result = ImageProductionResult(
        status="ready_for_delivery",
        production_ready=True,
        records=tuple(records),
        failures=(),
    )
    return bundle, manifest, package, result


class ComicV2DeliveryTests(unittest.TestCase):
    def test_delivery_blocks_when_any_planned_image_is_missing(self):
        from src.comic_office.v2.delivery import DeliveryValidationError, build_delivery_from_v2

        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            bundle, manifest, package, result = parts(root)
            incomplete = ImageProductionResult(
                status="ready_for_delivery",
                production_ready=True,
                records=result.records[:1],
                failures=(),
            )

            with self.assertRaisesRegex(DeliveryValidationError, "缺少已批准图片"):
                build_delivery_from_v2(bundle, manifest, package, incomplete, root / "out")

    def test_delivery_blocks_stale_image_version(self):
        from dataclasses import replace
        from src.comic_office.v2.delivery import DeliveryValidationError, build_delivery_from_v2

        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            bundle, manifest, package, result = parts(root)
            stale = replace(result.records[0], style_version=99)
            result = replace(result, records=(stale,) + result.records[1:])

            with self.assertRaisesRegex(DeliveryValidationError, "版本"):
                build_delivery_from_v2(bundle, manifest, package, result, root / "out")

    def test_delivery_embeds_every_planned_image_and_shot_prompt(self):
        from src.comic_office.v2.delivery import build_delivery_from_v2
        from src.comic_office.v2.production import image_production_result_from_dict

        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            bundle, manifest, package, result = parts(root)

            delivery = build_delivery_from_v2(bundle, manifest, package, result, root / "out")

            self.assertTrue(delivery.audit.handoff_ready)
            self.assertEqual(delivery.audit.embedded_images, len(result.records))
            doc = Document(delivery.path)
            text = "\n".join(
                [p.text for p in doc.paragraphs]
                + [cell.text for table in doc.tables for row in table.rows for cell in row.cells]
            )
            self.assertIn("three_view", text)
            self.assertIn("expression_sheet", text)
            self.assertIn("视频生成提示词", text)
            self.assertIn("负面提示词", text)
            self.assertIn("镜头验收标准", text)
            self.assertIn("平台执行备注", text)
            self.assertEqual(image_production_result_from_dict(result.to_dict()), result)

    def test_delivery_writes_machine_readable_handoff_manifest(self):
        from src.comic_office.v2.delivery import build_delivery_from_v2

        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            bundle, manifest, package, result = parts(root)

            delivery = build_delivery_from_v2(bundle, manifest, package, result, root / "out")
            handoff = json.loads(delivery.handoff_manifest_path.read_text(encoding="utf-8"))

            self.assertEqual(handoff["schema"], "comic_production_handoff_manifest_v3")
            self.assertEqual(handoff["schema_version"], 3)
            self.assertEqual(handoff["story"]["story_id"], bundle.creative.story_id)
            self.assertEqual(handoff["story"]["story_version"], bundle.creative.story_version)
            self.assertEqual(handoff["story"]["source_story"], bundle.creative.source_story)
            self.assertEqual(handoff["story"]["main_conflict"], bundle.creative.main_conflict)
            self.assertEqual(handoff["story"]["causal_chain"], list(bundle.creative.causal_chain))
            self.assertEqual(handoff["style"]["style_id"], bundle.visual.style_id)
            self.assertEqual(handoff["style"]["palette"], list(bundle.visual.palette))
            self.assertEqual(handoff["style"]["character_rules"], list(bundle.visual.character_rules))
            self.assertEqual(handoff["manifest"]["manifest_id"], manifest.manifest_id)
            self.assertEqual(handoff["word_canvas"]["filename"], delivery.path.name)
            self.assertEqual(len(handoff["assets"]), len(manifest.items))
            first_asset = handoff["assets"][0]
            self.assertEqual(first_asset["type_label"], "人物")
            self.assertEqual(first_asset["visual_locks"], ["靛青长袍"])
            self.assertEqual(first_asset["allowed_changes"], ["表情", "姿势"])
            self.assertEqual(first_asset["review_status"], "awaiting_user_review")
            self.assertEqual(first_asset["identity_baseline_image_id"], result.records[0].image_id)
            self.assertEqual(first_asset["identity_baseline_image_kind"], result.records[0].image_kind)
            self.assertEqual(
                first_asset["image_ids_by_kind"],
                {
                    record.image_kind: record.image_id
                    for record in result.records
                    if record.asset_id == first_asset["asset_id"]
                },
            )
            self.assertEqual(
                sorted(image["image_id"] for image in handoff["images"]),
                sorted(record.image_id for record in result.records),
            )
            first_image = next(image for image in handoff["images"] if image["image_kind"] == "three_view")
            self.assertIn("林昭 three_view", first_image["generator_prompt"])
            self.assertIn("禁止文字水印", first_image["negative_prompt"])
            self.assertIn("禁止现代服装", first_image["negative_prompt"])
            self.assertEqual(first_image["story_id"], bundle.creative.story_id)
            self.assertEqual(first_image["style_id"], bundle.visual.style_id)
            self.assertEqual(first_image["manifest_version"], manifest.version)
            self.assertEqual(first_image["review"], {"status": "pass"})
            self.assertEqual(handoff["shots"][0]["shot_id"], package.shots[0].shot_id)
            self.assertEqual(handoff["shots"][0]["reference_asset_ids"], list(package.shots[0].reference_asset_ids))
            self.assertEqual(
                handoff["shots"][0]["acceptance_criteria"],
                list(package.shots[0].acceptance_criteria),
            )
            self.assertEqual(handoff["shots"][0]["platform_note"], package.shots[0].platform_note)
            reference_image = handoff["shots"][0]["reference_images"][0]
            self.assertEqual(reference_image["asset_id"], manifest.items[0].asset_id)
            self.assertEqual(reference_image["image_id"], result.records[0].image_id)
            self.assertEqual(reference_image["image_kind"], result.records[0].image_kind)
            self.assertTrue(reference_image["file"].endswith(".png"))
            self.assertEqual(handoff["shots"][0]["first_frame_reference_image"], reference_image)
            director = handoff["shots"][0]["director_execution"]
            self.assertEqual(director["contract_version"], 1)
            self.assertEqual(director["style_id"], package.shots[0].style_id)
            self.assertEqual(director["style_version"], bundle.visual.style_version)
            self.assertEqual(director["first_frame_image_id"], reference_image["image_id"])
            self.assertEqual(director["reference_asset_ids"], list(package.shots[0].reference_asset_ids))
            self.assertEqual(director["action_chain"], list(package.shots[0].action_chain))
            self.assertEqual(director["performance_intent"], package.shots[0].performance_intent)
            self.assertEqual(director["framing"], package.shots[0].framing)
            self.assertEqual(director["camera_movement"], package.shots[0].camera_movement)
            self.assertEqual(director["lighting"], package.shots[0].lighting)
            self.assertEqual(director["dialogue"], package.shots[0].dialogue)
            self.assertEqual(director["sound"], package.shots[0].sound)
            reference_asset = handoff["shots"][0]["reference_asset_chain"][0]
            self.assertEqual(reference_asset["asset_id"], manifest.items[0].asset_id)
            self.assertEqual(reference_asset["name"], manifest.items[0].name)
            self.assertEqual(reference_asset["first_frame_file"], reference_image["file"])
            self.assertEqual(handoff["shots"][0]["video_prompt_block"], package.shots[0].generator_prompt)
            self.assertEqual(
                handoff["shots"][0]["negative_prompt_block"],
                "；".join(package.shots[0].negative_prompt),
            )
            self.assertIn("绑定首帧参考图片", handoff["shots"][0]["execution_steps"][0])
            self.assertIn("粘贴视频提示词", handoff["shots"][0]["execution_steps"][1])
            self.assertIn("按验收标准检查", handoff["shots"][0]["execution_steps"][2])
            lineage = handoff["production_lineage"]
            self.assertEqual(
                [stage["stage"] for stage in lineage],
                [
                    "story_contract",
                    "visual_bible",
                    "asset_manifest",
                    "prompt_package",
                    "image_production",
                    "visual_review",
                    "delivery",
                ],
            )
            for stage in lineage:
                self.assertTrue(stage["department"])
                self.assertTrue(stage["agent"])
                self.assertTrue(stage["status"])
                self.assertTrue(stage["human_checkpoint"])
                self.assertTrue(stage["handoff_to"])
                self.assertTrue(stage["acceptance_criteria"])
            self.assertTrue(handoff["audit"]["handoff_ready"])
            benchmark = handoff["quality_benchmark"]
            self.assertEqual(benchmark["benchmark"], "comic_production_package_quality")
            self.assertEqual(benchmark["benchmark_version"], 1)
            self.assertEqual(benchmark["status"], "needs_review")
            self.assertFalse(benchmark["production_quality_verified"])


if __name__ == "__main__":
    unittest.main()
