import base64
import tempfile
import unittest
from pathlib import Path

from docx import Document
from docx.shared import Inches

from src.comic_office.v2.asset_manifest import build_asset_manifest
from src.comic_office.v2.contracts import build_contract_bundle
from src.comic_office.v2.prompt_director import build_shot_card
from src.comic_office.v2.word_canvas import build_word_canvas_v2


PNG_1X1 = base64.b64decode(
    "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mP8/x8AAusB9Y9ZrWQAAAAASUVORK5CYII="
)


def delivery_parts():
    story = "林昭抱着裂纹月灯进入中央月塔。她逆向转动控制环，最终熄灭月塔。"
    bundle = build_contract_bundle(story, {
        "title": "借月人",
        "genre": "古风幻想",
        "theme": "记忆与光明的代价",
        "protagonist_goal": "熄灭月塔",
        "main_conflict": "月塔依靠记忆维持光明",
        "causal_chain": ["进入月塔", "转动控制环", "熄灭月塔"],
        "ending": "林昭最终熄灭月塔",
        "episodes": [{"episode": 1, "summary": "熄灭月塔", "evidence_quote": "林昭抱着裂纹月灯进入中央月塔"}],
        "visual": {
            "medium": "电影级国风厚涂动画",
            "era": "架空古代",
            "aspect_ratio": "9:16",
            "palette": ["靛青", "银白", "暗朱红"],
            "lighting": "冷银月光与暗红火光对照",
            "camera_language": "稳定构图，克制运镜",
            "character_rules": ["脸型固定"],
            "costume_rules": ["古代窄袖长袍"],
            "prop_rules": ["裂纹位置固定"],
            "architecture_rules": ["石塔与古铜结构"],
            "visual_motifs": ["裂纹月灯"],
            "prohibited_elements": ["现代车辆", "可读文字"],
        },
    })
    manifest = build_asset_manifest(bundle, [
        {
            "asset_type": "character", "name": "林昭", "evidence_quote": "林昭",
            "scene_ids": ["scene_01"], "story_purpose": "完成最终选择",
            "visual_locks": ["靛青窄袖长袍"], "allowed_changes": ["表情", "姿势"],
        },
        {
            "asset_type": "prop", "name": "裂纹月灯", "evidence_quote": "裂纹月灯",
            "scene_ids": ["scene_01"], "story_purpose": "触发最终选择",
            "visual_locks": ["固定弧形裂纹"], "allowed_changes": ["亮度"],
        },
        {
            "asset_type": "scene", "name": "中央月塔", "evidence_quote": "中央月塔",
            "scene_ids": ["scene_01"], "story_purpose": "最终选择发生地",
            "visual_locks": ["圆形外环", "中央控制环"], "allowed_changes": ["灯火状态"],
        },
    ])
    by_type = {item.asset_type: item for item in manifest.items}
    shot = build_shot_card({
        "shot_id": "SHOT-01",
        "scene_id": "scene_01",
        "story_beat": "林昭熄灭月塔",
        "action_chain": ["走向控制环", "双手握紧", "逆向转动", "灯火逐层熄灭"],
        "performance_intent": "克制而坚定",
        "framing": "35mm 中广角",
        "camera_movement": "缓慢跟随并绕到正面",
        "lighting": "冷银顶光逐层熄灭",
        "dialogue": "林昭：够了。",
        "sound": "控制环摩擦声后归于寂静",
        "retry_strategy": "调度不稳时只保留林昭与控制环。",
    }, characters=[by_type["character"]], props=[by_type["prop"]], scene=by_type["scene"], visual=bundle.visual)
    return bundle, manifest, (shot,)


def all_text(doc):
    return "\n".join(
        [paragraph.text for paragraph in doc.paragraphs]
        + [cell.text for table in doc.tables for row in table.rows for cell in row.cells]
    )


class ComicV2WordCanvasTests(unittest.TestCase):
    def test_canvas_uses_page_sections_not_nine_column_table(self):
        bundle, manifest, shots = delivery_parts()
        with tempfile.TemporaryDirectory() as tmp:
            image = Path(tmp) / "asset.png"
            image.write_bytes(PNG_1X1)
            images = {item.asset_id: str(image) for item in manifest.items}

            result = build_word_canvas_v2(bundle, manifest, shots, images, Path(tmp))
            doc = Document(result.path)

            self.assertFalse(any(len(table.columns) >= 9 for table in doc.tables))
            self.assertIn("视觉母版", all_text(doc))
            self.assertIn("SHOT-01", all_text(doc))
            self.assertEqual(result.audit.max_table_columns, 2)

    def test_each_asset_and_shot_has_a_dedicated_heading(self):
        bundle, manifest, shots = delivery_parts()
        with tempfile.TemporaryDirectory() as tmp:
            image = Path(tmp) / "asset.png"
            image.write_bytes(PNG_1X1)
            images = {item.asset_id: str(image) for item in manifest.items}

            result = build_word_canvas_v2(bundle, manifest, shots, images, Path(tmp))
            doc = Document(result.path)
            headings = "\n".join(
                paragraph.text for paragraph in doc.paragraphs
                if paragraph.style and paragraph.style.name.startswith("Heading")
            )

            for item in manifest.items:
                self.assertIn(item.asset_id, headings)
            self.assertIn("SHOT-01", headings)

    def test_audit_counts_images_and_reports_missing_references(self):
        bundle, manifest, shots = delivery_parts()
        with tempfile.TemporaryDirectory() as tmp:
            image = Path(tmp) / "asset.png"
            image.write_bytes(PNG_1X1)
            first = manifest.items[0]

            result = build_word_canvas_v2(bundle, manifest, shots, {first.asset_id: str(image)}, Path(tmp))

            self.assertEqual(result.audit.embedded_images, 1)
            self.assertEqual(result.audit.asset_count, 3)
            self.assertEqual(result.audit.shot_count, 1)
            self.assertEqual(set(result.audit.missing_image_asset_ids), {item.asset_id for item in manifest.items[1:]})
            self.assertFalse(result.audit.handoff_ready)

    def test_complete_structure_reopens_and_is_handoff_ready(self):
        bundle, manifest, shots = delivery_parts()
        with tempfile.TemporaryDirectory() as tmp:
            image = Path(tmp) / "asset.png"
            image.write_bytes(PNG_1X1)
            images = {item.asset_id: str(image) for item in manifest.items}

            result = build_word_canvas_v2(bundle, manifest, shots, images, Path(tmp))
            reopened = Document(result.path)

            self.assertGreater(len(reopened.paragraphs), 10)
            self.assertEqual(result.audit.missing_image_asset_ids, ())
            self.assertEqual(result.audit.structural_errors, ())
            self.assertTrue(result.audit.handoff_ready)

    def test_canvas_uses_compact_print_margins_for_asset_cards(self):
        bundle, manifest, shots = delivery_parts()
        with tempfile.TemporaryDirectory() as tmp:
            image = Path(tmp) / "asset.png"
            image.write_bytes(PNG_1X1)
            images = {item.asset_id: str(image) for item in manifest.items}

            result = build_word_canvas_v2(bundle, manifest, shots, images, Path(tmp))
            doc = Document(result.path)
            section = doc.sections[0]

            self.assertLessEqual(section.top_margin, Inches(0.8))
            self.assertLessEqual(section.bottom_margin, Inches(0.8))
            self.assertLessEqual(section.left_margin, Inches(0.8))
            self.assertLessEqual(section.right_margin, Inches(0.8))


if __name__ == "__main__":
    unittest.main()
