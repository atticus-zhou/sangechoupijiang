"""Word delivery canvas for the AI comic office."""

from __future__ import annotations

import re
from pathlib import Path


def build_comic_word_canvas(package: dict, image_artifacts: list[dict], output_dir: Path) -> Path:
    """Create a docx production canvas and return its path."""
    try:
        from docx import Document
        from docx.enum.section import WD_ORIENT
        from docx.shared import Inches, Pt
    except Exception as exc:  # pragma: no cover - depends on optional runtime package
        raise RuntimeError(f"python-docx is not available: {exc}") from exc

    output_dir.mkdir(parents=True, exist_ok=True)
    title = package.get("title") or "AI漫剧制片画布"
    path = output_dir / f"{_safe_filename(title)}_制片画布.docx"

    doc = Document()
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width
    section.left_margin = Inches(0.35)
    section.right_margin = Inches(0.35)
    section.top_margin = Inches(0.35)
    section.bottom_margin = Inches(0.35)

    doc.add_heading(f"{title} - AI漫剧制片画布", level=1)
    brief = package.get("creative_brief", {}) or {}
    intro = doc.add_paragraph()
    intro.add_run("核心灵感：").bold = True
    intro.add_run(str(brief.get("core_idea") or title))
    doc.add_paragraph(f"故事承诺：{brief.get('story_promise', '')}")
    doc.add_paragraph(f"用户确认/补充：{package.get('user_answers') or '暂无'}")
    script = package.get("confirmed_script", {}) or package.get("script_preview", {}) or {}
    doc.add_heading("完整故事稿", level=2)
    doc.add_paragraph(script.get("story_draft") or "当前任务未附带完整故事稿。")
    image_by_ref = _image_artifact_map(image_artifacts)
    global_negative = package.get("global_negative_prompt") or ""
    if global_negative:
        doc.add_heading("全局一致性与负面提示词", level=2)
        doc.add_paragraph("下面这组规则对全部人物、道具、场景、镜头画面提示词和视频提示词生效，不再在每个镜头里重复。")
        doc.add_paragraph(global_negative)

    _add_asset_tables(doc, package, image_by_ref)
    _add_shot_template_table(doc, package)

    table = doc.add_table(rows=1, cols=9)
    table.style = "Table Grid"
    headers = ["镜头", "画面内容", "参考资产", "动作链", "表演意图", "摄影", "灯光", "镜头导演提示词", "视频生成提示词"]
    for index, header in enumerate(headers):
        table.rows[0].cells[index].text = header

    for shot in package.get("shots", []) or []:
        cells = table.add_row().cells
        cells[0].text = shot.get("id", "")
        cells[1].text = shot.get("beat", "")
        cells[2].text = shot.get("identity_references") or shot.get("reference_assets", "") or _execution_asset_note(shot)
        cells[3].text = shot.get("action_chain", "")
        cells[4].text = shot.get("performance_intent", "")
        cells[5].text = shot.get("cinematography", "")
        cells[6].text = shot.get("lighting", "")
        cells[7].text = shot.get("director_prompt", "") or shot.get("image_prompt", "")
        cells[8].text = shot.get("video_prompt", "")

    doc.add_heading("平台执行表（Libtv / 图生视频）", level=2)
    doc.add_paragraph(
        "这一页给后续视频平台执行使用：复制视频提示词，并参考人物、道具、场景基础资产，保持角色、场景和道具锚点一致。"
    )
    exec_table = doc.add_table(rows=1, cols=7)
    exec_table.style = "Table Grid"
    exec_headers = ["镜头", "参考资产", "视频时长", "平台提示词", "镜头运动", "动作链", "失败重试建议"]
    for index, header in enumerate(exec_headers):
        exec_table.rows[0].cells[index].text = header

    for shot in package.get("shots", []) or []:
        cells = exec_table.add_row().cells
        shot_id = shot.get("id", "")
        cells[0].text = shot_id
        cells[1].text = shot.get("identity_references") or _execution_asset_note(shot)
        cells[2].text = _shot_duration(shot)
        cells[3].text = _platform_prompt(shot)
        cells[4].text = shot.get("camera_movement") or shot.get("camera") or "保持主体稳定，轻微推进。"
        cells[5].text = shot.get("action_chain", "") or shot.get("beat", "")
        cells[6].text = _retry_advice(shot)

    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            run.font.name = "Microsoft YaHei"
            run.font.size = Pt(9)
    for doc_table in doc.tables:
        for row in doc_table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.name = "Microsoft YaHei"
                        run.font.size = Pt(8)

    doc.save(path)
    return path


def _add_asset_tables(doc, package: dict, image_by_ref: dict[str, Path]) -> None:
    groups = [
        ("人物资产", package.get("characters", []) or [], ["ID", "名称", "职责", "设定图", "三视图/表情表提示词"]),
        ("道具资产", package.get("props", []) or [], ["ID", "名称", "连续性规则", "设定图", "多角度/使用状态提示词"]),
        ("场景资产", package.get("scenes", []) or [], ["ID", "名称", "连续性规则", "设定图", "空间/机位提示词"]),
    ]
    for heading, items, headers in groups:
        doc.add_heading(heading, level=2)
        if not items:
            doc.add_paragraph("暂无。")
            continue
        table = doc.add_table(rows=1, cols=len(headers))
        table.style = "Table Grid"
        for index, header in enumerate(headers):
            table.rows[0].cells[index].text = header
        for item in items:
            cells = table.add_row().cells
            cells[0].text = item.get("id", "")
            cells[1].text = item.get("name", "")
            cells[2].text = item.get("role") or item.get("continuity_rule") or item.get("visual_lock", "")
            cells[3].text = item.get("image_prompt", "")
            _append_asset_picture(cells[3], item.get("id", ""), image_by_ref)
            cells[4].text = _asset_specs_text(item)
            _append_asset_spec_pictures(cells[4], item, image_by_ref)


def _add_shot_template_table(doc, package: dict) -> None:
    shots = package.get("shots", []) or []
    if not shots:
        return
    doc.add_heading("镜头模板与执行说明", level=2)
    doc.add_paragraph(
        "这一页给导演/图生视频平台使用：先看镜头模板和构图，再复制后面的导演提示词。"
    )
    table = doc.add_table(rows=1, cols=7)
    table.style = "Table Grid"
    headers = ["镜头", "模板", "用途", "构图", "参考身份证", "平台执行说明", "失败重试优先级"]
    for index, header in enumerate(headers):
        table.rows[0].cells[index].text = header
    for shot in shots:
        cells = table.add_row().cells
        cells[0].text = shot.get("id", "")
        cells[1].text = shot.get("shot_template") or shot.get("framing", "")
        cells[2].text = shot.get("shot_template_purpose") or shot.get("beat", "")
        cells[3].text = shot.get("composition") or shot.get("cinematography", "")
        cells[4].text = shot.get("identity_references") or shot.get("reference_assets", "")
        cells[5].text = shot.get("platform_note") or shot.get("camera_movement", "")
        cells[6].text = _retry_advice(shot)


def _asset_specs_text(item: dict) -> str:
    lines = []
    card = item.get("identity_card") or {}
    if card:
        lines.append(f"{card.get('label', '资产身份证')}：{', '.join(card.get('image_refs') or [])}")
        if card.get("usage_rule"):
            lines.append(f"使用规则：{card.get('usage_rule')}")
    for spec in item.get("asset_specs", []) or []:
        lines.append(f"{spec.get('label', spec.get('kind', 'asset'))}：{spec.get('prompt', '')}")
        if spec.get("acceptance"):
            lines.append(f"验收：{spec.get('acceptance')}")
    return "\n".join(lines)


def _append_asset_spec_pictures(cell, item: dict, image_by_ref: dict[str, Path]) -> None:
    item_id = item.get("id", "")
    for spec in item.get("asset_specs", []) or []:
        source_id = f"{item_id}_{spec.get('kind', '')}".strip("_")
        _append_asset_picture(cell, source_id, image_by_ref, label=spec.get("label") or spec.get("kind", "资产图"))


def _append_asset_picture(cell, source_id: str, image_by_ref: dict[str, Path], label: str = "生成图") -> None:
    from docx.shared import Inches

    if not source_id:
        return
    image_path = image_by_ref.get(source_id)
    if not image_path:
        return
    paragraph = cell.add_paragraph()
    paragraph.add_run(f"{label}：{image_path.name}")
    try:
        paragraph.add_run("\n").add_picture(str(image_path), width=Inches(1.2))
    except Exception:
        paragraph.add_run(f"\n图片路径：{image_path}")


def _image_artifact_map(image_artifacts: list[dict]) -> dict[str, Path]:
    result = {}
    for artifact in image_artifacts or []:
        if artifact.get("artifact_type") != "generated_image":
            continue
        metadata = artifact.get("metadata") or {}
        source_id = metadata.get("source_id") or ""
        local_path = metadata.get("local_path") or metadata.get("path") or ""
        if source_id and local_path and Path(local_path).exists():
            result[source_id] = Path(local_path)
    return result


def _execution_asset_note(shot: dict) -> str:
    parts = []
    characters = ", ".join(shot.get("characters", []) or [])
    props = ", ".join(shot.get("props", []) or [])
    scene = shot.get("scene", "")
    if characters:
        parts.append(f"人物：{characters}")
    if scene:
        parts.append(f"场景：{scene}")
    if props:
        parts.append(f"道具：{props}")
    return "\n".join(parts) or "参考基础资产"


def _shot_duration(shot: dict) -> str:
    duration = shot.get("duration") or shot.get("video_duration") or ""
    if duration:
        return str(duration)
    return "5-6秒"


def _platform_prompt(shot: dict) -> str:
    parts = [
        shot.get("video_prompt", ""),
        f"导演提示词：{shot.get('director_prompt') or shot.get('image_prompt', '')}",
    ]
    return "\n".join(part for part in parts if part)


def _retry_advice(shot: dict) -> str:
    anchors = []
    characters = ", ".join(shot.get("characters", []) or [])
    props = ", ".join(shot.get("props", []) or [])
    scene = shot.get("scene", "")
    if characters:
        anchors.append(f"人物：{characters}")
    if scene:
        anchors.append(f"场景：{scene}")
    if props:
        anchors.append(f"道具：{props}")
    anchor_text = "；".join(anchors) or "先固定角色和场景锚点"
    return f"若脸、服装或场景跑偏，回到人物/道具/场景基础资产重新约束；重试时保留：{anchor_text}。"


def _safe_filename(text: str) -> str:
    clean = re.sub(r'[<>:"/\\|?*\s]+', "_", text).strip("._")
    return clean[:80] or "comic_canvas"
