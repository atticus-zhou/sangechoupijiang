"""Page-based Word production canvas for comic production V2."""

from __future__ import annotations

import re
import shutil
import subprocess
from dataclasses import dataclass
from pathlib import Path
from typing import Callable

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from .asset_manifest import AssetManifest, AssetPlan
from .contracts import ContractBundle
from .prompt_director import PromptPlan, ShotCard


INK = "243246"
INDIGO = "314C75"
SILVER = "E8EEF5"
PALE = "F4F6F9"
VERMILION = "A94A44"
MUTED = "687386"
CONTENT_WIDTH_DXA = 9360


@dataclass(frozen=True)
class VisualQASummary:
    status: str
    preview_page_images: tuple[str, ...]
    reason: str


@dataclass(frozen=True)
class DocumentAudit:
    embedded_images: int
    asset_count: int
    shot_count: int
    missing_image_asset_ids: tuple[str, ...]
    structural_errors: tuple[str, ...]
    max_table_columns: int
    handoff_ready: bool
    visual_qa_status: str = "skipped"
    preview_page_images: tuple[str, ...] = ()
    visual_qa_reason: str = "未执行页面渲染 QA"
    portfolio_ready: bool = False
    portfolio_notes: tuple[str, ...] = ()


@dataclass(frozen=True)
class CanvasBuildResult:
    path: Path
    audit: DocumentAudit
    handoff_manifest_path: Path | None = None


def build_word_canvas_v2(
    bundle: ContractBundle,
    manifest: AssetManifest,
    shots: tuple[ShotCard, ...] | list[ShotCard],
    image_paths: dict[str, str | dict[str, str]],
    output_dir: Path,
    *,
    asset_prompts: dict[tuple[str, str], PromptPlan] | None = None,
    require_all_planned_images: bool = False,
    render_visual_qa: bool = True,
    visual_qa_renderer: Callable[[Path, Path], VisualQASummary] | None = None,
) -> CanvasBuildResult:
    """Build and structurally audit a portrait, page-based production canvas."""
    output_dir.mkdir(parents=True, exist_ok=True)
    path = output_dir / f"{_safe_filename(bundle.creative.title)}_V2制片画布.docx"
    doc = Document()
    _configure_document(doc)
    _add_running_furniture(doc, bundle.creative.title)
    _add_cover(doc, bundle)
    _add_project_overview(doc, bundle, manifest, tuple(shots))
    _add_story(doc, bundle)
    _add_visual_bible(doc, bundle)

    prompt_lookup = asset_prompts or {}
    missing_images = []
    expected_image_count = 0
    for asset in manifest.items:
        resolved = _resolve_asset_images(asset, image_paths)
        expected_kinds = asset.planned_images if require_all_planned_images else asset.planned_images[:1]
        expected_image_count += len(expected_kinds)
        for image_kind in expected_kinds:
            image_path = resolved.get(image_kind)
            if not image_path or not image_path.exists():
                missing_images.append(
                    f"{asset.asset_id}/{image_kind}" if require_all_planned_images else asset.asset_id
                )
            _add_asset_page(
                doc,
                asset,
                image_path,
                image_kind=image_kind,
                prompt=prompt_lookup.get((asset.asset_id, image_kind)),
            )
    _add_continuity_page(doc, manifest)
    for shot in shots:
        _add_shot_page(doc, shot, manifest, image_paths)
    _add_handoff_page(doc, manifest, shots)
    doc.save(path)
    visual_qa = (
        (visual_qa_renderer or _render_docx_visual_qa)(path, output_dir / "preview")
        if render_visual_qa
        else VisualQASummary(status="skipped", preview_page_images=(), reason="未启用 Word 页面渲染 QA")
    )

    audit = _audit_document(
        path,
        manifest,
        tuple(shots),
        tuple(missing_images),
        expected_image_count=expected_image_count,
        visual_qa=visual_qa,
    )
    return CanvasBuildResult(path=path, audit=audit)


def _configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.75)
    section.right_margin = Inches(0.75)
    section.bottom_margin = Inches(0.72)
    section.left_margin = Inches(0.75)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)

    normal = doc.styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25
    for name, size, color, before, after in (
        ("Heading 1", 16, INDIGO, 18, 10),
        ("Heading 2", 13, INDIGO, 14, 7),
        ("Heading 3", 12, INK, 10, 5),
    ):
        style = doc.styles[name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True


def _add_running_furniture(doc: Document, title: str) -> None:
    section = doc.sections[0]
    header = section.header.paragraphs[0]
    header.text = f"三个臭皮匠 | {title}制片画布"
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in header.runs:
        _font(run, 8.5, MUTED)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = footer.add_run("第 ")
    _font(run, 8.5, MUTED)
    _add_field(run, "PAGE")
    tail = footer.add_run(" 页")
    _font(tail, 8.5, MUTED)


def _add_cover(doc: Document, bundle: ContractBundle) -> None:
    for _ in range(5):
        doc.add_paragraph()
    kicker = doc.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = kicker.add_run("AI COMIC PRODUCTION CANVAS")
    _font(run, 9, VERMILION, bold=True)
    kicker.paragraph_format.space_after = Pt(18)
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(bundle.creative.title)
    _font(run, 30, INDIGO, bold=True)
    title.paragraph_format.space_after = Pt(8)
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("独立制片画布 V2")
    _font(run, 15, MUTED)
    subtitle.paragraph_format.space_after = Pt(26)
    _add_callout(
        doc,
        "使用方式",
        "先锁定故事与视觉母版，再使用资产身份证生成镜头。每张镜头执行卡都引用已批准资产，并附带失败重试策略。",
        fill=PALE,
        accent=INDIGO,
    )
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta.add_run(
        f"{bundle.creative.genre}  |  {bundle.visual.aspect_ratio}  |  "
        f"Story v{bundle.creative.story_version}  |  Style v{bundle.visual.style_version}"
    )
    _font(run, 9, MUTED)


def _add_project_overview(
    doc: Document,
    bundle: ContractBundle,
    manifest: AssetManifest,
    shots: tuple[ShotCard, ...],
) -> None:
    _new_page_heading(doc, "1. 项目概览")
    _add_callout(
        doc,
        "交付摘要",
        (
            "这份画布用于把已确认故事、视觉母版、资产身份证、镜头执行卡和下游生产清单放在同一份文档中。"
            "下游可以按资产 ID 找参考图，按镜头卡执行视频生成。"
        ),
        fill=PALE,
        accent=INDIGO,
    )
    _add_label_detail(doc, [
        ("项目名称", bundle.creative.title),
        ("题材类型", bundle.creative.genre),
        ("核心命题", bundle.creative.theme),
        ("版本", f"Story v{bundle.creative.story_version} / Style v{bundle.visual.style_version}"),
        ("资产数量", str(len(manifest.items))),
        ("镜头数量", str(len(shots))),
        ("资产清单", f"Manifest v{manifest.version} / {manifest.manifest_hash[:12]}"),
        ("交付状态", "等待结构审计通过后交付" if manifest.items and shots else "缺少资产或镜头，不能交付"),
    ])
    _add_callout(
        doc,
        "下游读取顺序",
        "先看视觉母版，再查人物/道具/场景资产页，最后按镜头执行卡逐条生成视频片段。",
        fill=SILVER,
        accent=INDIGO,
        compact=True,
    )


def _add_story(doc: Document, bundle: ContractBundle) -> None:
    _new_page_heading(doc, "2. 完整故事")
    _add_label_detail(doc, [
        ("核心命题", bundle.creative.theme),
        ("人物目标", bundle.creative.protagonist_goal),
        ("主冲突", bundle.creative.main_conflict),
        ("结局", bundle.creative.ending),
    ])
    paragraph = doc.add_paragraph(bundle.creative.source_story)
    paragraph.paragraph_format.space_before = Pt(10)
    paragraph.paragraph_format.line_spacing = 1.35
    doc.add_heading("分集结构", level=2)
    for episode in bundle.creative.episodes:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        label = p.add_run(f"第 {episode.episode} 集  ")
        _font(label, 10.5, INDIGO, bold=True)
        _font(p.add_run(episode.summary), 10.5, INK)


def _add_visual_bible(doc: Document, bundle: ContractBundle) -> None:
    visual = bundle.visual
    _new_page_heading(doc, "3. 视觉母版")
    _add_callout(
        doc,
        "视觉承诺",
        f"{visual.medium}，{visual.era}。以{'、'.join(visual.palette)}为主色，{visual.lighting}。",
        fill=SILVER,
        accent=INDIGO,
    )
    _add_label_detail(doc, [
        ("风格身份证", visual.style_id),
        ("画面比例", visual.aspect_ratio),
        ("镜头语言", visual.camera_language),
        ("人物规则", "；".join(visual.character_rules)),
        ("服装规则", "；".join(visual.costume_rules)),
        ("道具规则", "；".join(visual.prop_rules)),
        ("建筑规则", "；".join(visual.architecture_rules)),
        ("视觉母题", "；".join(visual.visual_motifs)),
    ])
    _add_callout(
        doc,
        "项目级禁用元素",
        "；".join(f"禁止{item.removeprefix('禁止')}" for item in visual.prohibited_elements),
        fill="FAF4F4",
        accent=VERMILION,
    )


def _add_asset_page(
    doc: Document,
    asset: AssetPlan,
    image_path: Path | None,
    *,
    image_kind: str = "",
    prompt: PromptPlan | None = None,
) -> None:
    label = {"character": "人物", "prop": "道具", "scene": "场景"}[asset.asset_type]
    suffix = f" | {image_kind}" if image_kind else ""
    _new_page_heading(doc, f"{asset.asset_id} | {asset.name}{suffix}")
    p = doc.add_paragraph()
    run = p.add_run(f"{label}资产身份证")
    _font(run, 10, VERMILION, bold=True)
    if image_path and image_path.exists():
        picture = doc.add_paragraph()
        picture.alignment = WD_ALIGN_PARAGRAPH.CENTER
        image_width = {"character": 3.5, "prop": 3.2, "scene": 2.9}[asset.asset_type]
        picture.add_run().add_picture(str(image_path), width=Inches(image_width))
        caption = doc.add_paragraph(f"已批准基础资产：{image_path.name}")
        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in caption.runs:
            _font(run, 8.5, MUTED)
    else:
        _add_callout(doc, "图片待补", "此资产尚未绑定批准图片，不能进入最终交付。", fill="FAF4F4", accent=VERMILION)
    _add_label_detail(doc, [
        ("故事用途", asset.story_purpose),
        ("原文证据", asset.evidence.evidence_quote),
        ("出现场次", "、".join(asset.evidence.scene_ids)),
        ("固定项", "；".join(asset.visual_locks)),
        ("允许变化", "；".join(asset.allowed_changes)),
        ("本页图种", image_kind or "基础身份图"),
        ("生产角色", prompt.production_role if prompt else ""),
        ("干净背景要求", "是" if prompt and prompt.clean_background_required else "否"),
        ("批准图片文件", image_path.name if image_path else "未绑定"),
        ("计划图片", "、".join(asset.planned_images)),
    ], compact=True)
    if prompt is not None:
        doc.add_heading("本图生成提示词", level=2)
        paragraph = doc.add_paragraph(prompt.generator_prompt)
        paragraph.paragraph_format.line_spacing = 1.08
        for run in paragraph.runs:
            _font(run, 8.8, INK)
        doc.add_heading("负面提示词", level=3)
        negative = doc.add_paragraph("；".join(prompt.negative_prompt))
        for run in negative.runs:
            _font(run, 8.8, INK)


def _add_continuity_page(doc: Document, manifest: AssetManifest) -> None:
    _new_page_heading(doc, "4. 连续性状态")
    _add_callout(
        doc,
        "执行原则",
        "镜头可以改变构图、表情、动作和灯光状态，但不得改变资产身份证中列出的固定项。",
        fill=SILVER,
        accent=INDIGO,
    )
    rows = []
    for item in manifest.items:
        rows.append((f"{item.asset_id}\n{item.name}", "；".join(item.visual_locks)))
    _add_label_detail(doc, rows)


def _add_shot_page(
    doc: Document,
    shot: ShotCard,
    manifest: AssetManifest,
    image_paths: dict[str, str | dict[str, str]],
) -> None:
    _new_page_heading(doc, f"{shot.shot_id} | 镜头执行卡")
    _add_callout(doc, "叙事任务", shot.story_beat, fill=SILVER, accent=INDIGO, compact=True)
    reference_rows = _shot_reference_rows(shot, manifest, image_paths)
    if reference_rows:
        _add_label_detail(doc, [
            ("首帧参考图片", "；".join(row["image"] for row in reference_rows)),
            ("引用资产链路", "；".join(f"{row['asset_id']} / {row['name']} / {row['image']}" for row in reference_rows)),
        ], compact=True)
    _add_label_detail(doc, [
        ("参考资产", "、".join(shot.reference_asset_ids)),
        ("动作链", " -> ".join(shot.action_chain)),
        ("表演意图", shot.performance_intent),
        ("摄影", f"{shot.framing}；{shot.camera_movement}"),
        ("灯光", shot.lighting),
        ("台词", shot.dialogue),
        ("声音", shot.sound),
        ("平台执行备注", shot.platform_note),
    ], compact=True)
    _add_callout(
        doc,
        "视频平台执行步骤",
        "先上传并绑定首帧参考图片，再粘贴视频提示词；生成失败时先检查资产引用、动作链和运镜，再按失败重试策略降级。"
        f"平台备注：{shot.platform_note}",
        fill=PALE,
        accent=INDIGO,
        compact=True,
    )
    doc.add_heading("视频生成提示词 / 视频提示词复制区", level=2)
    prompt = doc.add_paragraph(shot.generator_prompt)
    prompt.paragraph_format.line_spacing = 1.08
    prompt.paragraph_format.space_after = Pt(3)
    for run in prompt.runs:
        _font(run, 8.8, INK)
    doc.add_heading("负面提示词 / 负面提示词复制区", level=3)
    negative = doc.add_paragraph("；".join(shot.negative_prompt))
    negative.paragraph_format.line_spacing = 1.05
    negative.paragraph_format.space_after = Pt(3)
    for run in negative.runs:
        _font(run, 8.8, INK)
    _add_callout(
        doc,
        "镜头验收标准",
        "；".join(shot.acceptance_criteria),
        fill=PALE,
        accent=INDIGO,
        compact=True,
    )
    _add_callout(doc, "失败重试", shot.retry_strategy, fill="FAF4F4", accent=VERMILION, compact=True)


def _shot_reference_rows(
    shot: ShotCard,
    manifest: AssetManifest,
    image_paths: dict[str, str | dict[str, str]],
) -> list[dict[str, str]]:
    asset_by_id = {item.asset_id: item for item in manifest.items}
    rows = []
    for asset_id in shot.reference_asset_ids:
        asset = asset_by_id.get(asset_id)
        if asset is None:
            rows.append({"asset_id": asset_id, "name": "未找到资产", "image": "未绑定图片"})
            continue
        resolved = _resolve_asset_images(asset, image_paths)
        first_kind = asset.planned_images[0] if asset.planned_images else ""
        image_path = resolved.get(first_kind) or next(iter(resolved.values()), None)
        rows.append({
            "asset_id": asset.asset_id,
            "name": asset.name,
            "image": image_path.name if image_path else "未绑定图片",
        })
    return rows


def _add_handoff_page(doc: Document, manifest: AssetManifest, shots: tuple[ShotCard, ...] | list[ShotCard]) -> None:
    _new_page_heading(doc, "5. 下游生产清单")
    _add_label_detail(doc, [
        ("资产版本", f"Manifest v{manifest.version} / {manifest.manifest_hash[:12]}"),
        ("资产数量", str(len(manifest.items))),
        ("镜头数量", str(len(shots))),
        ("组装顺序", " -> ".join(shot.shot_id for shot in shots)),
    ])
    _add_callout(
        doc,
        "完成标准",
        "所有人物、道具和场景均有批准图片；每张镜头卡引用真实资产；跨图质检通过；失败镜头按照卡片中的重试策略降级；handoff manifest 与 Word 画布引用一致。",
        fill=PALE,
        accent=INDIGO,
    )
    _add_agent_handoff_summary(doc)


def _add_agent_handoff_summary(doc: Document) -> None:
    doc.add_heading("多 Agent 交接与验收", level=2)
    _add_callout(
        doc,
        "阅读方式",
        "每一行代表一个办公室交接点。上游产物满足验收后，才交给下一步继续生产。",
        fill=SILVER,
        accent=INDIGO,
        compact=True,
    )
    _add_label_detail(doc, [
        ("故事合同", "交给：视觉母版；验收：故事原文、故事版本和禁止改写规则齐全。"),
        ("风格圣经", "交给：资产拆解；验收：视觉母版包含画风、时代、比例、色彩、服装和禁用元素。"),
        ("资产拆解", "交给：提示词与镜头执行包；验收：每个资产都有原文证据、故事用途、计划图片和审核状态。"),
        ("提示词与镜头执行包", "交给：基础图片生产；验收：资产提示词和镜头卡引用已审核资产，并把负面提示词单独列出。"),
        ("基础图片生产", "交给：一致性质检；验收：人物和道具基础图保持干净白底，场景图保留空间信息。"),
        ("一致性质检", "交给：Word 画布交付；验收：图片通过身份、风格、时代、空间和用途检查，风险项有处理结论。"),
        ("Word 画布交付", "交给：下游视频平台；验收：Word 画布、图片、镜头卡和 handoff manifest 可下载且引用一致。"),
    ], compact=True)


def _new_page_heading(doc: Document, text: str) -> None:
    heading = doc.add_heading(text, level=1)
    heading.paragraph_format.page_break_before = True


def _add_label_detail(doc: Document, rows: list[tuple[str, str]], *, compact: bool = False) -> None:
    table = doc.add_table(rows=0, cols=2)
    table.autofit = False
    _set_table_geometry(table, (1800, 7560))
    for label, detail in rows:
        cells = table.add_row().cells
        cells[0].text = label
        cells[1].text = detail
        cells[0].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        cells[1].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        _shade(cells[0], PALE)
        for cell in cells:
            cell.paragraphs[0].paragraph_format.space_after = Pt(0)
            cell.paragraphs[0].paragraph_format.line_spacing = 1.05 if compact else 1.15
        for run in cells[0].paragraphs[0].runs:
            _font(run, 8.4 if compact else 9, INDIGO, bold=True)
        for run in cells[1].paragraphs[0].runs:
            _font(run, 8.8 if compact else 9.5, INK)
    _set_table_geometry(table, (1800, 7560))
    vertical_margin = 55 if compact else 80
    _set_cell_margins(table, top=vertical_margin, bottom=vertical_margin, start=120, end=120)


def _add_callout(
    doc: Document,
    title: str,
    body: str,
    *,
    fill: str,
    accent: str,
    compact: bool = False,
) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.autofit = False
    _set_table_geometry(table, (CONTENT_WIDTH_DXA,))
    cell = table.cell(0, 0)
    _shade(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.05 if compact else 1.15
    _font(p.add_run(f"{title}\n"), 9.2 if compact else 10, accent, bold=True)
    _font(p.add_run(body), 8.8 if compact else 9.5, INK)
    vertical_margin = 80 if compact else 120
    _set_cell_margins(table, top=vertical_margin, bottom=vertical_margin, start=160, end=160)


def _audit_document(
    path: Path,
    manifest: AssetManifest,
    shots: tuple[ShotCard, ...],
    missing_images: tuple[str, ...],
    *,
    expected_image_count: int,
    visual_qa: VisualQASummary,
) -> DocumentAudit:
    doc = Document(path)
    asset_ids = {item.asset_id for item in manifest.items}
    errors = []
    if not manifest.items:
        errors.append("asset_manifest_empty")
    if not shots:
        errors.append("shot_cards_empty")
    for shot in shots:
        missing_refs = [ref for ref in shot.reference_asset_ids if ref not in asset_ids]
        if missing_refs:
            errors.append(f"{shot.shot_id}:missing_asset_refs:{','.join(missing_refs)}")
    max_columns = max((len(table.columns) for table in doc.tables), default=0)
    if max_columns > 2:
        errors.append("table_has_more_than_two_columns")
    embedded = len(doc.inline_shapes)
    text = "\n".join(
        [paragraph.text for paragraph in doc.paragraphs]
        + [cell.text for table in doc.tables for row in table.rows for cell in row.cells]
    )
    for shot in shots:
        printed_refs = [ref for ref in shot.reference_asset_ids if ref in text]
        if not printed_refs:
            errors.append(f"{shot.shot_id}:missing_printed_asset_refs")
    handoff_ready = not missing_images and not errors and embedded >= expected_image_count
    portfolio_notes = _portfolio_notes(doc, manifest, shots, handoff_ready, visual_qa)
    return DocumentAudit(
        embedded_images=embedded,
        asset_count=len(manifest.items),
        shot_count=len(shots),
        missing_image_asset_ids=missing_images,
        structural_errors=tuple(errors),
        max_table_columns=max_columns,
        handoff_ready=handoff_ready,
        visual_qa_status=visual_qa.status,
        preview_page_images=visual_qa.preview_page_images,
        visual_qa_reason=visual_qa.reason,
        portfolio_ready=handoff_ready and not any(note.startswith("缺少") for note in portfolio_notes),
        portfolio_notes=portfolio_notes,
    )


def _portfolio_notes(
    doc: Document,
    manifest: AssetManifest,
    shots: tuple[ShotCard, ...],
    handoff_ready: bool,
    visual_qa: VisualQASummary,
) -> tuple[str, ...]:
    text = "\n".join(
        [paragraph.text for paragraph in doc.paragraphs]
        + [cell.text for table in doc.tables for row in table.rows for cell in row.cells]
    )
    notes = []
    required_sections = (
        ("封面", "AI COMIC PRODUCTION CANVAS"),
        ("项目概览", "1. 项目概览"),
        ("完整故事", "2. 完整故事"),
        ("视觉母版", "3. 视觉母版"),
        ("资产页", next((item.asset_id for item in manifest.items), "")),
        ("镜头执行卡", next((shot.shot_id for shot in shots), "")),
        ("下游生产清单", "5. 下游生产清单"),
    )
    for label, marker in required_sections:
        if marker and marker in text:
            notes.append(f"{label}已就绪")
        else:
            notes.append(f"缺少{label}")
    if visual_qa.status == "ready":
        notes.append(f"页面预览已生成{len(visual_qa.preview_page_images)}张")
    elif visual_qa.status == "skipped":
        notes.append(f"页面预览待补：{visual_qa.reason}")
    else:
        notes.append(f"页面预览失败：{visual_qa.reason}")
    if handoff_ready:
        notes.append("交付结构审计已通过")
    return tuple(notes)


def _render_docx_visual_qa(path: Path, preview_dir: Path) -> VisualQASummary:
    office = shutil.which("soffice") or shutil.which("libreoffice")
    if not office:
        return VisualQASummary(
            status="skipped",
            preview_page_images=(),
            reason="未找到 soffice/libreoffice，当前环境无法把 Word 渲染成页面图片",
        )
    try:
        import fitz  # type: ignore
    except Exception:
        return VisualQASummary(
            status="skipped",
            preview_page_images=(),
            reason="未安装 PyMuPDF，当前环境无法把 PDF 页面渲染成 PNG",
        )
    preview_dir.mkdir(parents=True, exist_ok=True)
    pdf_dir = preview_dir / "pdf"
    pdf_dir.mkdir(parents=True, exist_ok=True)
    try:
        subprocess.run(
            [
                office,
                "--headless",
                "--convert-to",
                "pdf",
                "--outdir",
                str(pdf_dir),
                str(path),
            ],
            check=True,
            stdout=subprocess.PIPE,
            stderr=subprocess.PIPE,
            timeout=60,
        )
        pdf_path = pdf_dir / f"{path.stem}.pdf"
        if not pdf_path.exists():
            return VisualQASummary("failed", (), "LibreOffice 未生成 PDF 预览文件")
        doc = fitz.open(str(pdf_path))
        rendered = []
        for index, page_index in enumerate(range(min(3, doc.page_count)), start=1):
            page = doc.load_page(page_index)
            pix = page.get_pixmap(matrix=fitz.Matrix(1.2, 1.2), alpha=False)
            out = preview_dir / f"{path.stem}_page_{index:03d}.png"
            pix.save(str(out))
            rendered.append(str(out))
        doc.close()
        status = "ready" if rendered else "failed"
        reason = "页面渲染 QA 已完成" if rendered else "PDF 没有可渲染页面"
        return VisualQASummary(status=status, preview_page_images=tuple(rendered), reason=reason)
    except Exception as exc:
        return VisualQASummary(
            status="failed",
            preview_page_images=(),
            reason=f"Word 页面渲染失败：{exc}",
        )


def _resolve_asset_images(
    asset: AssetPlan,
    image_paths: dict[str, str | dict[str, str]],
) -> dict[str, Path]:
    raw = image_paths.get(asset.asset_id)
    if isinstance(raw, dict):
        return {
            str(kind): Path(path)
            for kind, path in raw.items()
            if str(path).strip()
        }
    if isinstance(raw, str) and raw.strip():
        return {asset.planned_images[0]: Path(raw)}
    resolved = {}
    for kind in asset.planned_images:
        value = image_paths.get(f"{asset.asset_id}:{kind}")
        if isinstance(value, str) and value.strip():
            resolved[kind] = Path(value)
    return resolved


def _set_table_geometry(table, widths: tuple[int, ...]) -> None:
    total = sum(widths)
    table.width = Inches(total / 1440)
    tbl_pr = table._tbl.tblPr
    for tag in ("w:tblW", "w:tblInd", "w:tblLayout"):
        existing = tbl_pr.find(qn(tag))
        if existing is not None:
            tbl_pr.remove(existing)
    tbl_w = OxmlElement("w:tblW")
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_pr.append(tbl_w)
    tbl_ind = OxmlElement("w:tblInd")
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_pr.append(tbl_ind)
    layout = OxmlElement("w:tblLayout")
    layout.set(qn("w:type"), "fixed")
    tbl_pr.append(layout)
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            cell.width = Inches(widths[index] / 1440)
            tc_w = cell._tc.get_or_add_tcPr().get_or_add_tcW()
            tc_w.set(qn("w:w"), str(widths[index]))
            tc_w.set(qn("w:type"), "dxa")


def _set_cell_margins(table, *, top: int, bottom: int, start: int, end: int) -> None:
    for row in table.rows:
        for cell in row.cells:
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_mar = tc_pr.first_child_found_in("w:tcMar")
            if tc_mar is None:
                tc_mar = OxmlElement("w:tcMar")
                tc_pr.append(tc_mar)
            for key, value in (("top", top), ("bottom", bottom), ("start", start), ("end", end)):
                node = tc_mar.find(qn(f"w:{key}"))
                if node is None:
                    node = OxmlElement(f"w:{key}")
                    tc_mar.append(node)
                node.set(qn("w:w"), str(value))
                node.set(qn("w:type"), "dxa")


def _shade(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), fill)


def _font(run, size: float, color: str, *, bold: bool = False) -> None:
    run.font.name = "Microsoft YaHei"
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.bold = bold


def _add_field(run, instruction: str) -> None:
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, end])


def _safe_filename(text: str) -> str:
    clean = re.sub(r'[<>:"/\\|?*\s]+', "_", text).strip("._")
    return clean[:80] or "comic_canvas"
