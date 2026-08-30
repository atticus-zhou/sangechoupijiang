"""Build and verify the deterministic comic-production V2 sample."""

from __future__ import annotations

import argparse
import hashlib
import json
import sys
from pathlib import Path
from typing import Any

from docx import Document
from PIL import Image, ImageDraw

REPO_ROOT = Path(__file__).resolve().parents[1]
if str(REPO_ROOT) not in sys.path:
    sys.path.insert(0, str(REPO_ROOT))

from src.comic_office.v2.asset_manifest import build_asset_manifest
from src.comic_office.v2.contracts import build_contract_bundle
from src.comic_office.v2.delivery import build_delivery_from_v2
from src.comic_office.v2.production import ImageProductionResult, ImageRecord, PromptPackage
from src.comic_office.v2.prompt_director import (
    PROMPT_STRATEGY_HASH,
    PROMPT_STRATEGY_VERSION,
    build_asset_prompt_plan,
    build_shot_card,
)
from src.comic_office.v2.visual_review import REVIEW_DIMENSIONS


def verify_delivery(fixture_path: Path, output_dir: Path) -> dict:
    fixture = json.loads(fixture_path.read_text(encoding="utf-8"))
    bundle = build_contract_bundle(fixture["source_story"], fixture["planner_payload"])
    manifest = build_asset_manifest(bundle, fixture["assets"])
    assets_by_name = {item.name: item for item in manifest.items}
    shots = []
    for payload in fixture["shots"]:
        shots.append(build_shot_card(
            payload,
            characters=[assets_by_name[name] for name in payload["characters"]],
            props=[assets_by_name[name] for name in payload["props"]],
            scene=assets_by_name[payload["scene"]],
            visual=bundle.visual,
        ))

    image_dir = output_dir / "fixture-images"
    image_dir.mkdir(parents=True, exist_ok=True)
    prompts = []
    records = []
    for asset in manifest.items:
        for index, image_kind in enumerate(asset.planned_images):
            prompt = build_asset_prompt_plan(asset, bundle.visual, image_kind=image_kind)
            prompts.append(prompt)
            image_path = image_dir / f"{asset.asset_id}_{image_kind}.png"
            _write_placeholder_image(image_path, asset.name, image_kind, asset.asset_type)
            records.append(ImageRecord(
                image_id=f"img_{asset.asset_id}_{image_kind}",
                asset_id=asset.asset_id,
                image_kind=image_kind,
                prompt_hash=hashlib.sha256(prompt.render().encode("utf-8")).hexdigest(),
                path=str(image_path),
                provider="fixture",
                model="deterministic-placeholder",
                attempts=1,
                status="approved",
                is_identity_baseline=index == 0,
                reference_image_ids=() if index == 0 else (f"img_{asset.asset_id}_{asset.planned_images[0]}",),
                story_id=bundle.creative.story_id,
                story_version=bundle.creative.story_version,
                style_id=bundle.visual.style_id,
                style_version=bundle.visual.style_version,
                manifest_version=manifest.version,
                review={
                    "status": "pass",
                    "handoff_ready": True,
                    "fixture": True,
                    "issues": [],
                    "failed_dimensions": [],
                    "missing_dimensions": [],
                    "scores": _fixture_visual_scores(asset_type=asset.asset_type, image_kind=image_kind),
                    "evidence": [
                        "fixture image generated from deterministic placeholder",
                        "asset id, style id, story version, and reference chain are preserved",
                        "fixture record is for structure and handoff demonstration only",
                    ],
                    "consistency_status": "pass",
                    "recovery_action": "",
                    "recovery_reason": "",
                    "rework_label": "",
                    "operator_steps": [],
                    "summary": "无 Key 固定样例，仅验证流程、引用和交付结构。",
                },
                production_role=prompt.production_role,
                clean_background_required=prompt.clean_background_required,
            ))
    prompt_package = PromptPackage(
        package_id="prompts_fixture",
        story_id=bundle.creative.story_id,
        story_version=bundle.creative.story_version,
        style_id=bundle.visual.style_id,
        style_version=bundle.visual.style_version,
        manifest_id=manifest.manifest_id,
        manifest_version=manifest.version,
        prompts=tuple(prompts),
        shots=tuple(shots),
    )
    image_result = ImageProductionResult(
        status="ready_for_delivery",
        production_ready=True,
        records=tuple(records),
        failures=(),
    )
    build = build_delivery_from_v2(bundle, manifest, prompt_package, image_result, output_dir)
    doc = Document(build.path)
    text = "\n".join(
        [paragraph.text for paragraph in doc.paragraphs]
        + [cell.text for table in doc.tables for row in table.rows for cell in row.cells]
    )
    missing_ids = [
        identifier
        for identifier in [*(item.asset_id for item in manifest.items), *(shot.shot_id for shot in shots)]
        if identifier not in text
    ]
    if missing_ids:
        raise AssertionError(f"delivery is missing IDs: {', '.join(missing_ids)}")
    word_canvas_agent_handoff = all(
        marker in text
        for marker in (
            "多 Agent 交接与验收",
            "交给",
            "验收",
            "资产拆解",
            "基础图片生产",
            "handoff manifest",
        )
    )
    if not word_canvas_agent_handoff:
        raise AssertionError("Word canvas is missing the multi-agent handoff checklist")
    handoff_manifest_path = build.handoff_manifest_path
    if handoff_manifest_path is None or not handoff_manifest_path.exists():
        raise AssertionError("delivery handoff manifest was not created")
    handoff_manifest = json.loads(handoff_manifest_path.read_text(encoding="utf-8"))
    if handoff_manifest.get("word_canvas", {}).get("filename") != build.path.name:
        raise AssertionError("handoff manifest does not point at the generated Word canvas")
    expected_image_files = [
        image.get("file")
        for image in (handoff_manifest.get("images") or [])
        if image.get("file")
    ]
    word_canvas_asset_file_references = (
        "批准图片文件" in text
        and bool(expected_image_files)
        and all(filename in text for filename in expected_image_files)
    )
    if not word_canvas_asset_file_references:
        raise AssertionError("Word canvas is missing approved asset image file references")
    image_prompt_ready = all(
        bool(image.get("generator_prompt")) and isinstance(image.get("negative_prompt"), list)
        for image in (handoff_manifest.get("images") or [])
    )
    if not image_prompt_ready:
        raise AssertionError("handoff manifest image records are missing executable prompts")
    image_prompt_strategy_ready = all(
        image.get("prompt_strategy_version") == PROMPT_STRATEGY_VERSION
        and image.get("prompt_strategy_hash") == PROMPT_STRATEGY_HASH
        for image in (handoff_manifest.get("images") or [])
    )
    if not image_prompt_strategy_ready:
        raise AssertionError("handoff manifest image records are missing prompt strategy lineage")
    image_production_roles_ready = all(
        bool(image.get("production_role"))
        and isinstance(image.get("clean_background_required"), bool)
        for image in (handoff_manifest.get("images") or [])
    )
    if not image_production_roles_ready:
        raise AssertionError("handoff manifest image records are missing production roles")
    asset_identity_ready = all(
        asset.get("type_label")
        and isinstance(asset.get("visual_locks"), list)
        and isinstance(asset.get("allowed_changes"), list)
        and bool(asset.get("review_status"))
        for asset in (handoff_manifest.get("assets") or [])
    )
    if not asset_identity_ready:
        raise AssertionError("handoff manifest asset records are missing identity fields")
    assets_by_id = {
        asset.get("asset_id"): asset
        for asset in (handoff_manifest.get("assets") or [])
        if asset.get("asset_id")
    }
    asset_baseline_chain_ready = all(
        bool(asset.get("identity_baseline_image_id"))
        and bool(asset.get("identity_baseline_image_kind"))
        and isinstance(asset.get("image_ids_by_kind"), dict)
        and asset.get("identity_baseline_image_id") in set(asset.get("image_ids") or [])
        and asset.get("identity_baseline_image_id") in set((asset.get("image_ids_by_kind") or {}).values())
        for asset in assets_by_id.values()
    )
    if not asset_baseline_chain_ready:
        raise AssertionError("handoff manifest asset records are missing baseline image chains")
    shot_reference_images_ready = all(
        isinstance(shot.get("reference_images"), list)
        and shot.get("reference_images")
        and all(item.get("asset_id") and item.get("image_id") and item.get("file") for item in shot.get("reference_images"))
        for shot in (handoff_manifest.get("shots") or [])
    )
    shot_reference_chain_ready = all(
        all(
            ref.get("image_id") in set((assets_by_id.get(ref.get("asset_id")) or {}).get("image_ids") or [])
            for ref in (shot.get("reference_images") or [])
        )
        for shot in (handoff_manifest.get("shots") or [])
    )
    shot_reference_images_ready = shot_reference_images_ready and shot_reference_chain_ready
    if not shot_reference_images_ready:
        raise AssertionError("handoff manifest shot records are missing reference images")
    shot_execution_notes_ready = all(
        isinstance(shot.get("acceptance_criteria"), list)
        and bool(shot.get("acceptance_criteria"))
        and bool(shot.get("platform_note"))
        for shot in (handoff_manifest.get("shots") or [])
    )
    if not shot_execution_notes_ready:
        raise AssertionError("handoff manifest shot records are missing execution notes")
    shot_production_package_ready = all(
        isinstance(shot.get("first_frame_reference_image"), dict)
        and bool((shot.get("first_frame_reference_image") or {}).get("file"))
        and isinstance(shot.get("reference_asset_chain"), list)
        and bool(shot.get("reference_asset_chain"))
        and all(
            item.get("asset_id")
            and item.get("name")
            and item.get("first_frame_file")
            for item in (shot.get("reference_asset_chain") or [])
        )
        and bool(shot.get("video_prompt_block"))
        and bool(shot.get("negative_prompt_block"))
        and isinstance(shot.get("execution_steps"), list)
        and len(shot.get("execution_steps") or []) >= 3
        for shot in (handoff_manifest.get("shots") or [])
    )
    if not shot_production_package_ready:
        raise AssertionError("handoff manifest shot records are missing production-ready shot packages")
    shot_prompt_strategy_ready = all(
        shot.get("prompt_strategy_version") == PROMPT_STRATEGY_VERSION
        and shot.get("prompt_strategy_hash") == PROMPT_STRATEGY_HASH
        and (shot.get("director_execution") or {}).get("prompt_strategy_version") == PROMPT_STRATEGY_VERSION
        for shot in (handoff_manifest.get("shots") or [])
    )
    if not shot_prompt_strategy_ready:
        raise AssertionError("handoff manifest shot records are missing prompt strategy lineage")
    asset_usage_map = handoff_manifest.get("asset_usage_map") or []
    asset_usage_ready = (
        isinstance(asset_usage_map, list)
        and len(asset_usage_map) == len(handoff_manifest.get("assets") or [])
        and all(
            item.get("asset_id")
            and item.get("name")
            and item.get("identity_baseline_image_id")
            and item.get("image_roles")
            and item.get("referenced_by_shots")
            and item.get("downstream_instruction")
            and item.get("handoff_ready") is True
            for item in asset_usage_map
        )
    )
    if not asset_usage_ready:
        raise AssertionError("handoff manifest is missing asset usage map")
    lineage = handoff_manifest.get("production_lineage") or []
    lineage_handoff_ready = (
        isinstance(lineage, list)
        and bool(lineage)
        and all(
            bool(item.get("handoff_to"))
            and bool(item.get("acceptance_criteria"))
            for item in lineage
        )
    )
    lineage_ready = (
        isinstance(lineage, list)
        and len(lineage) >= 6
        and all(
            item.get("stage")
            and item.get("department")
            and item.get("agent")
            and item.get("status")
            and item.get("human_checkpoint")
            for item in lineage
        )
        and lineage_handoff_ready
    )
    if not lineage_ready:
        raise AssertionError("handoff manifest is missing production lineage")
    prompt_package = handoff_manifest.get("prompt_package") or {}
    prompt_strategy_ready = (
        prompt_package.get("prompt_strategy_version") == PROMPT_STRATEGY_VERSION
        and prompt_package.get("prompt_strategy_hash") == PROMPT_STRATEGY_HASH
        and any(
            item.get("stage") == "prompt_package"
            and item.get("prompt_strategy_version") == PROMPT_STRATEGY_VERSION
            for item in lineage
        )
        and image_prompt_strategy_ready
        and shot_prompt_strategy_ready
    )
    if not prompt_strategy_ready:
        raise AssertionError("handoff manifest is missing prompt strategy lineage")
    quick_start = handoff_manifest.get("downstream_quick_start") or []
    shot_ids = {
        shot.get("shot_id")
        for shot in (handoff_manifest.get("shots") or [])
        if shot.get("shot_id")
    }
    quick_start_ready = (
        isinstance(quick_start, list)
        and len(quick_start) >= 5
        and [item.get("step") for item in quick_start] == list(range(1, len(quick_start) + 1))
        and all(
            item.get("title")
            and item.get("owner")
            and item.get("input_refs")
            and item.get("action")
            and item.get("output")
            and item.get("acceptance")
            for item in quick_start
        )
        and shot_ids.issubset(set(next(
            (
                item.get("input_refs") or []
                for item in quick_start
                if "镜头" in str(item.get("title") or "")
            ),
            [],
        )))
    )
    if not quick_start_ready:
        raise AssertionError("handoff manifest is missing the downstream quick-start playbook")
    audit = build.audit
    result = {
        "path": str(build.path),
        "handoff_manifest_path": str(handoff_manifest_path),
        "output_dir": str(output_dir),
        "word_canvas_exists": build.path.is_file(),
        "word_canvas_bytes": build.path.stat().st_size if build.path.is_file() else 0,
        "handoff_manifest_exists": True,
        "handoff_manifest_bytes": handoff_manifest_path.stat().st_size if handoff_manifest_path.is_file() else 0,
        "handoff_manifest_assets": len(handoff_manifest.get("assets") or []),
        "handoff_manifest_images": len(handoff_manifest.get("images") or []),
        "handoff_manifest_shots": len(handoff_manifest.get("shots") or []),
        "handoff_manifest_image_prompts": image_prompt_ready,
        "handoff_manifest_prompt_strategy": prompt_strategy_ready,
        "handoff_manifest_image_production_roles": image_production_roles_ready,
        "handoff_manifest_asset_identity_fields": asset_identity_ready,
        "handoff_manifest_asset_baseline_chain": asset_baseline_chain_ready,
        "handoff_manifest_shot_reference_images": shot_reference_images_ready,
        "handoff_manifest_shot_execution_notes": shot_execution_notes_ready,
        "handoff_manifest_shot_production_package": shot_production_package_ready,
        "handoff_manifest_asset_usage_map": asset_usage_ready,
        "handoff_manifest_asset_usage_map_items": len(asset_usage_map),
        "handoff_manifest_production_lineage": lineage_ready,
        "handoff_manifest_lineage_handoff_fields": lineage_handoff_ready,
        "handoff_manifest_downstream_quick_start": quick_start_ready,
        "handoff_manifest_downstream_quick_start_steps": len(quick_start),
        "word_canvas_agent_handoff": word_canvas_agent_handoff,
        "word_canvas_asset_file_references": word_canvas_asset_file_references,
        "handoff_ready": audit.handoff_ready,
        "asset_count": audit.asset_count,
        "shot_count": audit.shot_count,
        "embedded_images": audit.embedded_images,
        "max_table_columns": audit.max_table_columns,
        "missing_image_asset_ids": list(audit.missing_image_asset_ids),
        "structural_errors": list(audit.structural_errors),
    }
    if not result["handoff_ready"]:
        raise AssertionError(json.dumps(result, ensure_ascii=False, indent=2))
    return result


def _fixture_visual_scores(*, asset_type: str, image_kind: str) -> dict[str, int]:
    scores = {
        "identity_consistency": 92,
        "style_consistency": 90,
        "era_media": 90,
        "spatial_structure": 88,
        "asset_purity": 90,
        "anatomy": 88,
        "purpose_fit": 92,
    }
    if asset_type in {"character", "prop"}:
        scores["asset_purity"] = 96
        scores["spatial_structure"] = 86 if image_kind in {"three_view", "turnaround"} else 88
    if asset_type == "scene":
        scores["spatial_structure"] = 94
        scores["asset_purity"] = 86
    return {dimension: scores.get(dimension, 88) for dimension in REVIEW_DIMENSIONS}


def _write_placeholder_image(path: Path, name: str, image_kind: str, asset_type: str) -> None:
    palette = {
        "character": (229, 236, 245),
        "prop": (244, 240, 232),
        "scene": (225, 232, 228),
    }
    image = Image.new("RGB", (1200, 900), palette.get(asset_type, (235, 235, 235)))
    draw = ImageDraw.Draw(image)
    draw.rectangle((36, 36, 1164, 864), outline=(49, 76, 117), width=6)
    draw.text((80, 90), f"{name}\n{image_kind}\nV2 DELIVERY FIXTURE", fill=(36, 50, 70), spacing=18)
    image.save(path, format="PNG")


def format_markdown(result: dict[str, Any]) -> str:
    word_canvas_status = _file_status(result.get("word_canvas_exists"), result.get("word_canvas_bytes"))
    manifest_status = _file_status(result.get("handoff_manifest_exists"), result.get("handoff_manifest_bytes"))
    lines = [
        "# Comic V2 Delivery Audit",
        "",
        f"Status: `{'passed' if result.get('handoff_ready') else 'failed'}`",
        f"Output directory: `{_display_path(result.get('output_dir'))}`",
        f"Word canvas: {word_canvas_status}",
        f"Handoff manifest: {manifest_status}",
        "",
        "## Delivery Counts",
        "",
        f"- Assets: {result.get('asset_count')}",
        f"- Shots: {result.get('shot_count')}",
        f"- Embedded images: {result.get('embedded_images')}",
        f"- Handoff assets: {result.get('handoff_manifest_assets')}",
        f"- Handoff images: {result.get('handoff_manifest_images')}",
        f"- Handoff shots: {result.get('handoff_manifest_shots')}",
        f"- Downstream quick-start steps: {result.get('handoff_manifest_downstream_quick_start_steps')}",
        "",
        "## Quality Gates",
        "",
        "| Gate | Status |",
        "| --- | --- |",
    ]
    gates = [
        ("handoff_ready", "Overall handoff ready"),
        ("handoff_manifest_exists", "Machine-readable handoff manifest"),
        ("handoff_manifest_image_prompts", "Executable image prompts"),
        ("handoff_manifest_prompt_strategy", "Prompt strategy lineage"),
        ("handoff_manifest_image_production_roles", "Image production roles"),
        ("handoff_manifest_asset_identity_fields", "Asset identity fields"),
        ("handoff_manifest_asset_baseline_chain", "Asset baseline reference chain"),
        ("handoff_manifest_shot_reference_images", "Shot reference images"),
        ("handoff_manifest_shot_execution_notes", "Shot execution notes"),
        ("handoff_manifest_shot_production_package", "Shot production package"),
        ("handoff_manifest_asset_usage_map", "Asset usage map"),
        ("handoff_manifest_production_lineage", "Production lineage"),
        ("handoff_manifest_lineage_handoff_fields", "Lineage handoff fields"),
        ("handoff_manifest_downstream_quick_start", "Downstream quick-start playbook"),
        ("word_canvas_agent_handoff", "Word canvas agent handoff checklist"),
        ("word_canvas_asset_file_references", "Word canvas approved image file references"),
    ]
    for key, label in gates:
        lines.append(f"| {label} | {bool(result.get(key))} |")

    if result.get("structural_errors") or result.get("missing_image_asset_ids"):
        lines.extend(["", "## Issues", ""])
        for item in result.get("structural_errors") or []:
            lines.append(f"- {item}")
        if result.get("missing_image_asset_ids"):
            lines.append("- Missing image assets: " + ", ".join(result["missing_image_asset_ids"]))
    return "\n".join(lines) + "\n"


def _display_path(value: object) -> str:
    if not value:
        return ""
    path = Path(str(value))
    try:
        return path.resolve().relative_to(REPO_ROOT).as_posix()
    except ValueError:
        return path.name or str(value)


def _file_status(exists: object, size: object) -> str:
    if not exists:
        return "`missing`"
    return f"`present` ({int(size or 0)} bytes)"


def main() -> int:
    if hasattr(sys.stdout, "reconfigure"):
        sys.stdout.reconfigure(encoding="utf-8")
    parser = argparse.ArgumentParser()
    parser.add_argument(
        "--fixture",
        type=Path,
        default=Path("tests/fixtures/comic_v2_sample.json"),
    )
    parser.add_argument(
        "--output-dir",
        type=Path,
        default=Path("output/comic_v2_verification"),
    )
    parser.add_argument("--format", choices=["text", "json", "markdown"], default="text")
    args = parser.parse_args()
    result = verify_delivery(args.fixture, args.output_dir)
    if args.format == "json":
        print(json.dumps(result, ensure_ascii=False, indent=2))
    elif args.format == "markdown":
        print(format_markdown(result))
    else:
        print(f"V2 delivery verified: {result['path']}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
